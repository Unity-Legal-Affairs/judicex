from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

from .models import MatterDocument, MatterFact


_MONTHS_IT = {
    "gennaio": 1,
    "febbraio": 2,
    "marzo": 3,
    "aprile": 4,
    "maggio": 5,
    "giugno": 6,
    "luglio": 7,
    "agosto": 8,
    "settembre": 9,
    "ottobre": 10,
    "novembre": 11,
    "dicembre": 12,
}

_NUMBER_WORDS = {
    "uno": 1,
    "una": 1,
    "due": 2,
    "tre": 3,
    "quattro": 4,
    "cinque": 5,
    "sei": 6,
    "sette": 7,
    "otto": 8,
    "nove": 9,
    "dieci": 10,
    "quindici": 15,
    "venti": 20,
    "trenta": 30,
    "quaranta": 40,
    "cinquanta": 50,
    "sessanta": 60,
    "novanta": 90,
    "centoventi": 120,
}


TEXT_SUFFIXES = {"", ".txt", ".md", ".markdown", ".csv", ".json"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}
IMAGE_SUFFIXES = {".bmp", ".gif", ".heic", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}


@dataclass(slots=True)
class PrivateDocumentIngestion:
    content: str
    metadata: dict[str, Any] = field(default_factory=dict)
    suggested_kind: str = "document"


def sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def make_matter_id(title: str, client_name: str = "", area: str = "") -> str:
    base = "|".join([title.strip().lower(), client_name.strip().lower(), area.strip().lower()])
    digest = hashlib.sha1(base.encode("utf-8")).hexdigest()[:16]
    slug = re.sub(r"[^a-z0-9]+", "_", title.lower()).strip("_")[:48] or "matter"
    return f"matter:{slug}:{digest}"


def make_matter_document_id(matter_id: str, title: str, content: str) -> str:
    digest = hashlib.sha1(f"{matter_id}|{title}|{sha256_text(content)}".encode("utf-8")).hexdigest()[:20]
    return f"matterdoc:{digest}"


def read_private_text_file(path: str | Path) -> str:
    return read_private_document_file(path).content


def read_private_document_file(path: str | Path) -> PrivateDocumentIngestion:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return PrivateDocumentIngestion(
            content=file_path.read_text(encoding="utf-8"),
            metadata={"file_format": suffix.lstrip(".") or "text", "text_extraction": "plain_text"},
            suggested_kind="document",
        )
    if suffix in PDF_SUFFIXES:
        return _read_pdf_file(file_path)
    if suffix in DOCX_SUFFIXES:
        return _read_docx_file(file_path)
    if suffix in IMAGE_SUFFIXES:
        return _read_image_file(file_path)
    return _read_binary_file(file_path)


def _read_pdf_file(path: Path) -> PrivateDocumentIngestion:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise ValueError("PDF ingestion requires the pypdf package.") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[pagina {index}]\n{text.strip()}")
    content = "\n\n".join(pages).strip()
    status = "text_extracted" if content else "no_extractable_text"
    if not content:
        content = (
            f"PDF caricato: {path.name}.\n"
            "Il file non contiene testo estraibile con il parser PDF disponibile."
        )
    return PrivateDocumentIngestion(
        content=content,
        metadata={
            "file_format": "pdf",
            "page_count": len(reader.pages),
            "text_extraction": status,
            "extractor": "pypdf",
        },
        suggested_kind="pdf",
    )


def _read_docx_file(path: Path) -> PrivateDocumentIngestion:
    try:
        from docx import Document as DocxDocument
    except ModuleNotFoundError as exc:
        raise ValueError("DOCX ingestion requires the python-docx package.") from exc

    doc = DocxDocument(str(path))
    parts: list[str] = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    content = "\n".join(parts).strip()
    status = "text_extracted" if content else "no_extractable_text"
    if not content:
        content = f"DOCX caricato: {path.name}.\nIl file non contiene testo estraibile."
    return PrivateDocumentIngestion(
        content=content,
        metadata={
            "file_format": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "text_extraction": status,
            "extractor": "python-docx",
        },
        suggested_kind="docx",
    )


def _read_image_file(path: Path) -> PrivateDocumentIngestion:
    try:
        from PIL import Image
    except ModuleNotFoundError as exc:
        raise ValueError("Image ingestion requires the Pillow package.") from exc

    with Image.open(path) as image:
        metadata = {
            "file_format": (image.format or path.suffix.lstrip(".")).lower(),
            "width": image.width,
            "height": image.height,
            "mode": image.mode,
            "text_extraction": "not_applicable_without_ocr",
            "extractor": "Pillow",
        }
    return PrivateDocumentIngestion(
        content=(
            f"Immagine caricata: {path.name}.\n"
            f"Formato: {metadata['file_format']}; dimensioni: {metadata['width']}x{metadata['height']}."
        ),
        metadata=metadata,
        suggested_kind="image",
    )


def _read_binary_file(path: Path) -> PrivateDocumentIngestion:
    suffix = path.suffix.lower().lstrip(".") or "binary"
    return PrivateDocumentIngestion(
        content=(
            f"File caricato: {path.name}.\n"
            "Il formato è stato conservato come allegato, ma non contiene testo estraibile dal parser corrente."
        ),
        metadata={"file_format": suffix, "text_extraction": "unsupported_binary"},
        suggested_kind="attachment",
    )


def extract_matter_facts(document: MatterDocument) -> list[MatterFact]:
    facts: list[MatterFact] = []
    seen: set[tuple[str, str, str, str]] = set()
    for sentence_index, sentence in enumerate(_sentences(document.content)):
        for item in _date_facts(sentence):
            facts.append(_fact(document, sentence, sentence_index, len(facts), seen=seen, **item))
        for item in _amount_facts(sentence):
            facts.append(_fact(document, sentence, sentence_index, len(facts), seen=seen, **item))
        for item in _deadline_facts(sentence):
            facts.append(_fact(document, sentence, sentence_index, len(facts), seen=seen, **item))
        for item in _party_facts(sentence):
            facts.append(_fact(document, sentence, sentence_index, len(facts), seen=seen, **item))
    return [fact for fact in facts if fact.id]


def _fact(
    document: MatterDocument,
    sentence: str,
    sentence_index: int,
    fact_index: int,
    *,
    seen: set[tuple[str, str, str, str]],
    fact_type: str,
    label: str,
    text: str,
    value: str = "",
    unit: str = "",
    date_value: str = "",
    confidence: float = 0.85,
) -> MatterFact:
    key = (fact_type, label, value or date_value, text)
    if key in seen:
        return MatterFact(
            id="",
            matter_id=document.matter_id,
            document_id=document.id,
            fact_type=fact_type,
            label=label,
            text=text,
        )
    seen.add(key)
    stable = "|".join([document.id, str(sentence_index), fact_type, label, value, date_value, text])
    fact_id = f"matterfact:{hashlib.sha1(stable.encode('utf-8')).hexdigest()[:20]}"
    return MatterFact(
        id=fact_id,
        matter_id=document.matter_id,
        document_id=document.id,
        fact_type=fact_type,
        label=label,
        text=text,
        value=value,
        unit=unit,
        date_value=date_value,
        confidence=confidence,
        source_quote=sentence,
        metadata={"sentence_index": sentence_index, "fact_index": fact_index},
    )


def _sentences(text: str) -> list[str]:
    normalized = text.replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    chunks = re.split(r"(?<=[.;!?])\s+|\n+", normalized)
    out: list[str] = []
    for chunk in chunks:
        cleaned = chunk.strip(" \t\n")
        if len(cleaned) >= 4:
            out.append(cleaned)
    return out


def _date_facts(sentence: str) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    for match in re.finditer(r"\b([0-3]?\d)[/-]([01]?\d)[/-](20\d{2}|19\d{2})\b", sentence):
        day, month, year = (int(match.group(1)), int(match.group(2)), int(match.group(3)))
        iso = _safe_date(year, month, day)
        if iso:
            facts.append(
                {
                    "fact_type": "date",
                    "label": _date_label(sentence),
                    "text": match.group(0),
                    "date_value": iso,
                    "confidence": 0.95,
                }
            )
    for match in re.finditer(
        r"\b([0-3]?\d)\s+(" + "|".join(_MONTHS_IT.keys()) + r")\s+(20\d{2}|19\d{2})\b",
        sentence,
        flags=re.IGNORECASE,
    ):
        day = int(match.group(1))
        month = _MONTHS_IT[match.group(2).lower()]
        year = int(match.group(3))
        iso = _safe_date(year, month, day)
        if iso:
            facts.append(
                {
                    "fact_type": "date",
                    "label": _date_label(sentence),
                    "text": match.group(0),
                    "date_value": iso,
                    "confidence": 0.95,
                }
            )
    for match in re.finditer(r"\b(20\d{2}|19\d{2})-([01]\d)-([0-3]\d)\b", sentence):
        year, month, day = (int(match.group(1)), int(match.group(2)), int(match.group(3)))
        iso = _safe_date(year, month, day)
        if iso:
            facts.append(
                {
                    "fact_type": "date",
                    "label": _date_label(sentence),
                    "text": match.group(0),
                    "date_value": iso,
                    "confidence": 0.95,
                }
            )
    return facts


def _amount_facts(sentence: str) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    pattern = r"(?:€\s*|euro\s+)([0-9]{1,3}(?:\.[0-9]{3})*(?:,[0-9]{2})?|[0-9]+(?:,[0-9]{2})?)"
    for match in re.finditer(pattern, sentence, flags=re.IGNORECASE):
        value = _normalize_amount(match.group(1))
        facts.append(
            {
                "fact_type": "amount",
                "label": _amount_label(sentence),
                "text": match.group(0),
                "value": value,
                "unit": "EUR",
                "confidence": 0.95,
            }
        )
    return facts


def _deadline_facts(sentence: str) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    if not re.search(r"\b(entro|termine|scadenza|decorsi|non oltre)\b", sentence, flags=re.IGNORECASE):
        return facts
    pattern = r"\b([0-9]+|[A-Za-zÀ-ÿ]+)\s+(giorni|giorno|mesi|mese|anni|anno)\b"
    for match in re.finditer(pattern, sentence, flags=re.IGNORECASE):
        raw = match.group(1).lower()
        value = raw if raw.isdigit() else str(_NUMBER_WORDS.get(raw, ""))
        if not value:
            continue
        unit = _normalize_unit(match.group(2))
        facts.append(
            {
                "fact_type": "deadline",
                "label": _deadline_label(sentence),
                "text": match.group(0),
                "value": value,
                "unit": unit,
                "confidence": 0.9,
            }
        )
    return facts


def _party_facts(sentence: str) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    pattern = r"\b(creditore|debitore|cliente|controparte|ricorrente|resistente|attore|convenuto)\s*:\s*([^;\n]{2,120})"
    for match in re.finditer(pattern, sentence, flags=re.IGNORECASE):
        role = match.group(1).lower()
        value = _clean_party(match.group(2))
        if value:
            facts.append(
                {
                    "fact_type": "party",
                    "label": role,
                    "text": value,
                    "value": value,
                    "confidence": 0.9,
                }
            )
    match = re.search(r"\btra\s+([^.;]{2,120}?)\s+e\s+([^.;]{2,120})", sentence, flags=re.IGNORECASE)
    if match:
        left = _clean_party(match.group(1))
        right = _clean_party(match.group(2))
        if left:
            facts.append({"fact_type": "party", "label": "parte", "text": left, "value": left, "confidence": 0.75})
        if right:
            facts.append({"fact_type": "party", "label": "parte", "text": right, "value": right, "confidence": 0.75})
    return facts


def _safe_date(year: int, month: int, day: int) -> str:
    try:
        return date(year, month, day).isoformat()
    except ValueError:
        return ""


def _normalize_amount(raw: str) -> str:
    compact = raw.replace(".", "").replace(",", ".")
    if "." not in compact:
        compact += ".00"
    integer, decimal = compact.split(".", 1)
    decimal = (decimal + "00")[:2]
    return f"{int(integer)}.{decimal}"


def _normalize_unit(raw: str) -> str:
    lowered = raw.lower()
    if lowered.startswith("giorn"):
        return "giorni"
    if lowered.startswith("mes"):
        return "mesi"
    if lowered.startswith("ann"):
        return "anni"
    return lowered


def _date_label(sentence: str) -> str:
    lowered = sentence.lower()
    if "fattura" in lowered:
        return "data fattura"
    if "contratto" in lowered:
        return "data contratto"
    if "notifica" in lowered:
        return "data notifica"
    if "scaden" in lowered:
        return "data scadenza"
    return "data"


def _amount_label(sentence: str) -> str:
    lowered = sentence.lower()
    if "fattura" in lowered:
        return "importo fattura"
    if "credito" in lowered:
        return "importo credito"
    if "capitale" in lowered:
        return "capitale"
    return "importo"


def _deadline_label(sentence: str) -> str:
    lowered = sentence.lower()
    if "opposizione" in lowered:
        return "termine opposizione"
    if "pag" in lowered:
        return "termine pagamento"
    if "notifica" in lowered:
        return "termine notifica"
    return "termine"


def _clean_party(raw: str) -> str:
    value = raw.strip(" ,;:.")
    value = re.sub(r"\s+", " ", value)
    return value[:120]
