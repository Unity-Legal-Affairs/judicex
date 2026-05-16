from __future__ import annotations

import argparse
import base64
import csv
import hashlib
import html
import hmac
import io
import json
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from flask import Flask, Response, jsonify, render_template, request, send_file, session

from .agent_runtime import JudicexAgentRuntime
from .drafter import DraftingError, draft_atto, list_templates, load_template
from .llm_provider import (
    list_provider_models,
    load_dotenv,
    make_client,
    public_settings,
    resolve_settings,
    save_settings,
    test_provider,
)
from .official_sources import list_official_bundles, sync_official_bundle
from .store import LegalMemoryStore
from .tools import LegalMemoryTools


def _safe_positive_int(value: Any, *, default: int) -> int:
    try:
        number = int(value)
    except (TypeError, ValueError):
        return default
    return number if number > 0 else default


def create_app(*, db_path: str | Path, default_model: str = "", default_area: str = "", ollama_host: str = "http://127.0.0.1:11434") -> Flask:
    load_dotenv()
    app = Flask(__name__)
    app.secret_key = os.getenv("JUDICEX_SECRET_KEY", "").strip() or base64.urlsafe_b64encode(os.urandom(32)).decode("ascii")
    app.config.update(
        JUDICEX_DB_PATH=str(db_path),
        JUDICEX_DEFAULT_MODEL=default_model,
        JUDICEX_DEFAULT_AREA=default_area,
        JUDICEX_OLLAMA_HOST=ollama_host,
        JUDICEX_OCR_MODEL=os.getenv("JUDICEX_OCR_MODEL", "glm-5:cloud").strip() or "glm-5:cloud",
        JUDICEX_OCR_MAX_PDF_PAGES=_safe_positive_int(os.getenv("JUDICEX_OCR_MAX_PDF_PAGES", "20"), default=20),
        JUDICEX_LOCAL_PASSWORD=os.getenv("JUDICEX_LOCAL_PASSWORD", "").strip(),
        JUDICEX_LOCAL_PASSWORD_HASH=os.getenv("JUDICEX_LOCAL_PASSWORD_HASH", "").strip(),
        MAX_CONTENT_LENGTH=50 * 1024 * 1024,
    )

    def store() -> LegalMemoryStore:
        return LegalMemoryStore(app.config["JUDICEX_DB_PATH"])

    @app.before_request
    def require_local_auth():
        if _is_auth_open_path(request.path):
            return None
        auth = _local_auth_state(app, store)
        if not auth["configured"]:
            return None
        if session.get("judicex_authenticated") is True:
            return None
        if request.path.startswith("/api/"):
            return jsonify({"error": "Password locale richiesta.", "auth_required": True}), 401
        return Response(_login_page(), mimetype="text/html")

    @app.get("/")
    @app.get("/chat")
    @app.get("/onboarding")
    @app.get("/search")
    @app.get("/dashboard")
    @app.get("/documents")
    @app.get("/workflows")
    @app.get("/tables")
    @app.get("/drafts")
    @app.get("/tools")
    @app.get("/settings")
    @app.get("/provider-ai")
    @app.get("/memory")
    @app.get("/sources")
    @app.get("/security")
    @app.get("/backup")
    @app.get("/matters")
    def index() -> str:
        return render_template(
            "index.html",
            default_model=app.config["JUDICEX_DEFAULT_MODEL"],
            default_area=app.config["JUDICEX_DEFAULT_AREA"],
            ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
        )

    @app.get("/api/state")
    def api_state():
        with store() as memory:
            llm_settings = resolve_settings(
                memory,
                default_model=app.config["JUDICEX_DEFAULT_MODEL"],
                ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
            )
            return jsonify(
                {
                    "health": memory.health(),
                    "matters": memory.list_matters(top_k=100),
                    "areas": memory.list_areas(),
                    "agent_memories": memory.list_agent_memories(top_k=50),
                    "official_bundles": list_official_bundles(memory),
                    "onboarding": memory.get_app_setting("onboarding") or {"completed": False},
                    "workflow_packs": memory.list_workflow_packs(),
                    "llm": public_settings(llm_settings),
                    "defaults": {
                        "model": llm_settings["model"],
                        "area": app.config["JUDICEX_DEFAULT_AREA"],
                        "host": llm_settings["base_url"],
                        "provider": llm_settings["provider"],
                    },
                }
            )

    @app.get("/api/auth/status")
    def api_auth_status():
        auth = _local_auth_state(app, store)
        return jsonify(
            {
                "configured": bool(auth["configured"]),
                "authenticated": bool(session.get("judicex_authenticated")) if auth["configured"] else True,
                "source": auth["source"],
            }
        )

    @app.post("/api/auth/login")
    def api_auth_login():
        payload = _json_payload()
        password = str(payload.get("password", ""))
        auth = _local_auth_state(app, store)
        if not auth["configured"]:
            session["judicex_authenticated"] = True
            return jsonify({"status": "ok", "configured": False})
        if not _verify_local_password(password, auth):
            return _error("Password non corretta.", status=401)
        session["judicex_authenticated"] = True
        return jsonify({"status": "ok", "configured": True})

    @app.post("/api/auth/logout")
    def api_auth_logout():
        session.pop("judicex_authenticated", None)
        return jsonify({"status": "ok"})

    @app.post("/api/auth/setup")
    def api_auth_setup():
        auth = _local_auth_state(app, store)
        if auth["configured"] and session.get("judicex_authenticated") is not True:
            return _error("Autenticazione richiesta.", status=401)
        payload = _json_payload()
        password = str(payload.get("password", ""))
        if len(password) < 8:
            return _error("Usa una password di almeno 8 caratteri.", status=400)
        with store() as memory:
            memory.set_app_setting("local_auth", {"password_hash": _hash_password(password), "enabled": True})
        session["judicex_authenticated"] = True
        return jsonify({"status": "ok"})

    @app.get("/api/chat-sessions")
    def api_list_chat_sessions():
        status = request.args.get("status", "open")
        with store() as memory:
            return jsonify({"sessions": memory.list_chat_sessions(status=status or None, top_k=100)})

    @app.post("/api/chat-sessions")
    def api_create_chat_session():
        payload = _json_payload()
        matter_id = str(payload.get("matter_id", "")).strip()
        with store() as memory:
            if matter_id and memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            session = memory.create_chat_session(
                title=str(payload.get("title", "")).strip(),
                matter_id=matter_id,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            return jsonify({"session": session, "sessions": memory.list_chat_sessions(top_k=100)})

    @app.get("/api/chat-sessions/<path:session_id>")
    def api_get_chat_session(session_id: str):
        with store() as memory:
            session = memory.get_chat_session(session_id, full=True)
            if session is None:
                return _error(f"Chat non trovata: {session_id}", status=404)
            return jsonify({"session": session})

    @app.patch("/api/chat-sessions/<path:session_id>")
    def api_update_chat_session(session_id: str):
        payload = _json_payload()
        with store() as memory:
            matter_id = str(payload["matter_id"]).strip() if "matter_id" in payload else None
            if matter_id and memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            session = memory.update_chat_session(
                session_id,
                title=str(payload["title"]).strip() if "title" in payload else None,
                matter_id=matter_id,
                status=str(payload["status"]).strip() if "status" in payload else None,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            if session is None:
                return _error(f"Chat non trovata: {session_id}", status=404)
            return jsonify({"session": session, "sessions": memory.list_chat_sessions(top_k=100)})

    @app.delete("/api/chat-sessions/<path:session_id>")
    def api_delete_chat_session(session_id: str):
        with store() as memory:
            if not memory.delete_chat_session(session_id):
                return _error(f"Chat non trovata: {session_id}", status=404)
            return jsonify({"status": "deleted", "sessions": memory.list_chat_sessions(top_k=100)})

    @app.post("/api/chat-sessions/<path:session_id>/messages")
    def api_add_chat_session_message(session_id: str):
        payload = _json_payload()
        role = str(payload.get("role", "")).strip()
        content = str(payload.get("content", "")).strip()
        if not content:
            return _error("Messaggio vuoto.", status=400)
        with store() as memory:
            try:
                message = memory.add_chat_message(
                    session_id,
                    role,
                    content,
                    metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
                )
            except ValueError as exc:
                return _error(str(exc), status=404 if "non trovata" in str(exc).lower() else 400)
            session = memory.get_chat_session(session_id, full=True)
            return jsonify({"message": message, "session": session, "sessions": memory.list_chat_sessions(top_k=100)})

    @app.get("/api/search")
    def api_search():
        query = str(request.args.get("q", "")).strip()
        scope = str(request.args.get("scope", "all")).strip() or "all"
        matter_id = str(request.args.get("matter_id", "")).strip()
        if scope not in {"all", "chat", "matters", "documents", "facts", "memory"}:
            scope = "all"
        try:
            top_k = int(request.args.get("top_k", 12))
        except ValueError:
            top_k = 12
        top_k = max(1, min(top_k, 50))

        def excerpt(value: str, *, max_len: int = 260) -> str:
            text = re.sub(r"\s+", " ", str(value or "")).strip()
            if len(text) <= max_len:
                return text
            if query:
                pos = text.lower().find(query.lower())
                if pos >= 0:
                    start = max(0, pos - 80)
                    end = min(len(text), start + max_len)
                    prefix = "..." if start else ""
                    suffix = "..." if end < len(text) else ""
                    return f"{prefix}{text[start:end]}{suffix}"
            return text[:max_len] + "..."

        with store() as memory:
            matter_lookup = {item["id"]: item for item in memory.list_matters(top_k=100)}

            def matter_title(item_matter_id: str) -> str:
                matter = matter_lookup.get(item_matter_id)
                return matter["title"] if matter else ""

            results: list[dict[str, Any]] = []
            if scope in {"all", "chat"}:
                for session in memory.search_chat_sessions(query, top_k=top_k):
                    results.append(
                        {
                            "type": "chat",
                            "id": session["id"],
                            "session_id": session["id"],
                            "matter_id": session.get("matter_id", ""),
                            "title": session.get("title") or "Chat",
                            "subtitle": f"{session.get('message_count', 0)} messaggi"
                            + (f" · {matter_title(session.get('matter_id', ''))}" if session.get("matter_id") else ""),
                            "excerpt": excerpt(session.get("title", "")),
                            "updated_at": session.get("last_message_at") or session.get("updated_at", ""),
                        }
                    )
                for message in memory.search_chat_messages(query, top_k=top_k):
                    results.append(
                        {
                            "type": "message",
                            "id": message["id"],
                            "session_id": message["session_id"],
                            "matter_id": message.get("matter_id", ""),
                            "title": message.get("session_title") or "Messaggio chat",
                            "subtitle": f"{message.get('role', 'messaggio')}"
                            + (f" · {matter_title(message.get('matter_id', ''))}" if message.get("matter_id") else ""),
                            "excerpt": excerpt(message.get("content", "")),
                            "updated_at": message.get("created_at", ""),
                        }
                    )

            if scope in {"all", "matters"}:
                for matter in memory.search_matters(query, top_k=top_k):
                    results.append(
                        {
                            "type": "matter",
                            "id": matter["id"],
                            "matter_id": matter["id"],
                            "title": matter["title"],
                            "subtitle": " · ".join(part for part in [matter.get("client_name", ""), matter.get("area", "")] if part),
                            "excerpt": excerpt(matter.get("summary", "")),
                            "updated_at": matter.get("updated_at", ""),
                        }
                    )

            if scope in {"all", "documents"}:
                for document in memory.search_matter_documents(query, matter_id=matter_id or None, top_k=top_k):
                    results.append(
                        {
                            "type": "document",
                            "id": document["id"],
                            "document_id": document["id"],
                            "matter_id": document["matter_id"],
                            "title": document.get("title") or "Documento",
                            "subtitle": " · ".join(part for part in [document.get("kind", ""), matter_title(document["matter_id"])] if part),
                            "excerpt": excerpt(document.get("excerpt", "")),
                            "updated_at": document.get("updated_at", ""),
                        }
                    )

            if scope in {"all", "facts"}:
                for fact in memory.search_matter_facts(query, matter_id=matter_id or None, top_k=top_k):
                    results.append(
                        {
                            "type": "fact",
                            "id": fact["id"],
                            "matter_id": fact["matter_id"],
                            "document_id": fact.get("document_id", ""),
                            "title": fact.get("label") or fact.get("fact_type") or "Fatto",
                            "subtitle": " · ".join(part for part in [fact.get("fact_type", ""), matter_title(fact["matter_id"])] if part),
                            "excerpt": excerpt(fact.get("text", "") or fact.get("source_quote", "")),
                            "updated_at": fact.get("created_at", ""),
                        }
                    )

            if scope in {"all", "memory"}:
                for item in memory.search_agent_memories(query, matter_id=matter_id or None, top_k=top_k):
                    results.append(
                        {
                            "type": "memory",
                            "id": item["id"],
                            "matter_id": item.get("matter_id", ""),
                            "title": item.get("title") or "Memoria agente",
                            "subtitle": " · ".join(part for part in [item.get("kind", ""), item.get("scope", "")] if part),
                            "excerpt": excerpt(item.get("excerpt", "")),
                            "updated_at": item.get("updated_at", ""),
                        }
                    )

            results.sort(key=lambda item: str(item.get("updated_at") or ""), reverse=True)
            return jsonify({"query": query, "scope": scope, "results": results[: top_k * 4]})

    @app.get("/api/agent-memory")
    def api_list_agent_memory():
        query = str(request.args.get("q", "")).strip()
        matter_id = request.args.get("matter_id")
        kind = request.args.get("kind")
        scope = request.args.get("scope")
        top_k = _safe_positive_int(request.args.get("top_k"), default=50)
        with store() as memory:
            return jsonify(
                {
                    "memories": memory.search_agent_memories(
                        query,
                        kind=kind,
                        scope=scope,
                        matter_id=matter_id,
                        top_k=top_k,
                        full=False,
                    )
                }
            )

    @app.post("/api/agent-memory")
    def api_add_agent_memory():
        payload = _json_payload()
        title = str(payload.get("title", "")).strip()
        content = str(payload.get("content", "")).strip()
        if not title or not content:
            return _error("Titolo e contenuto memoria sono obbligatori.", status=400)
        tags = payload.get("tags") if isinstance(payload.get("tags"), list) else []
        with store() as memory:
            try:
                item = memory.add_agent_memory(
                    kind=str(payload.get("kind", "note")),
                    title=title,
                    content=content,
                    scope=str(payload.get("scope", "global")),
                    matter_id=str(payload.get("matter_id", "")),
                    tags=[str(tag) for tag in tags],
                    importance=float(payload.get("importance", 0.5)),
                    source=str(payload.get("source", "web")),
                    metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
                )
            except ValueError as exc:
                return _error(str(exc), status=400)
            return jsonify({"memory": item, "memories": memory.list_agent_memories(top_k=50)})

    @app.delete("/api/agent-memory/<path:memory_id>")
    def api_delete_agent_memory(memory_id: str):
        with store() as memory:
            if not memory.delete_agent_memory(memory_id):
                return _error(f"Memoria non trovata: {memory_id}", status=404)
            return jsonify({"status": "deleted", "memories": memory.list_agent_memories(top_k=50)})

    @app.patch("/api/agent-memory/<path:memory_id>")
    def api_update_agent_memory(memory_id: str):
        payload = _json_payload()
        with store() as memory:
            current = memory.get_agent_memory(memory_id, full=True)
            if current is None:
                return _error(f"Memoria non trovata: {memory_id}", status=404)
            tags = payload.get("tags") if isinstance(payload.get("tags"), list) else current.get("tags", [])
            try:
                item = memory.add_agent_memory(
                    memory_id=memory_id,
                    kind=str(payload.get("kind", current["kind"])),
                    title=str(payload.get("title", current["title"])),
                    content=str(payload.get("content", current.get("content", ""))),
                    scope=str(payload.get("scope", current["scope"])),
                    matter_id=str(payload.get("matter_id", current.get("matter_id", ""))),
                    tags=[str(tag) for tag in tags],
                    importance=float(payload.get("importance", current.get("importance", 0.5))),
                    source=str(payload.get("source", current.get("source", "web"))),
                    metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else current.get("metadata", {}),
                )
            except ValueError as exc:
                return _error(str(exc), status=400)
            return jsonify({"memory": item, "memories": memory.list_agent_memories(top_k=50)})

    @app.post("/api/onboarding/complete")
    def api_complete_onboarding():
        with store() as memory:
            setting = memory.set_app_setting(
                "onboarding",
                {
                    "completed": True,
                    "completed_at": datetime.now(timezone.utc).isoformat(),
                },
            )
            return jsonify({"onboarding": setting})

    @app.get("/api/official-bundles")
    def api_official_bundles():
        with store() as memory:
            return jsonify({"bundles": list_official_bundles(memory)})

    @app.post("/api/official-bundles/<path:bundle_name>/sync")
    def api_sync_official_bundle(bundle_name: str):
        payload = _json_payload()
        with store() as memory:
            try:
                result = sync_official_bundle(
                    memory,
                    bundle_name=bundle_name,
                    as_of_date=str(payload.get("as_of_date", "")).strip() or None,
                )
            except Exception as exc:
                return _error(str(exc), status=400)
            return jsonify(
                {
                    "result": result,
                    "health": memory.health(),
                    "bundles": list_official_bundles(memory),
                }
            )

    @app.get("/api/backup")
    def api_backup():
        return _backup_response(Path(app.config["JUDICEX_DB_PATH"]))

    @app.post("/api/restore")
    def api_restore():
        upload = request.files.get("file")
        if upload is None:
            return _error("Carica un file backup .zip.", status=400)
        try:
            result = _restore_backup(Path(app.config["JUDICEX_DB_PATH"]), upload)
        except Exception as exc:
            return _error(str(exc), status=400)
        session.pop("judicex_authenticated", None)
        return jsonify(result)

    @app.post("/api/security/password")
    def api_security_password():
        payload = _json_payload()
        password = str(payload.get("password", ""))
        if len(password) < 8:
            return _error("Usa una password di almeno 8 caratteri.", status=400)
        with store() as memory:
            memory.set_app_setting("local_auth", {"password_hash": _hash_password(password), "enabled": True})
        session["judicex_authenticated"] = True
        return jsonify({"status": "ok", "auth": {"configured": True, "authenticated": True}})

    @app.get("/api/settings/llm")
    def api_get_llm_settings():
        with store() as memory:
            settings = resolve_settings(
                memory,
                default_model=app.config["JUDICEX_DEFAULT_MODEL"],
                ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
            )
            return jsonify({"settings": public_settings(settings)})

    @app.get("/api/settings/llm/models")
    def api_list_llm_models():
        payload = {
            "provider": request.args.get("provider", ""),
            "model": request.args.get("model", ""),
            "base_url": request.args.get("base_url", ""),
        }
        with store() as memory:
            settings = resolve_settings(
                memory,
                default_model=app.config["JUDICEX_DEFAULT_MODEL"],
                ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
                overrides=payload,
            )
            return jsonify(list_provider_models(settings))

    @app.patch("/api/settings/llm")
    def api_update_llm_settings():
        payload = _json_payload()
        with store() as memory:
            settings = save_settings(
                memory,
                payload,
                default_model=app.config["JUDICEX_DEFAULT_MODEL"],
                ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
            )
            return jsonify({"settings": public_settings(settings)})

    @app.post("/api/settings/llm/test")
    def api_test_llm_settings():
        payload = _json_payload()
        with store() as memory:
            settings = resolve_settings(
                memory,
                default_model=app.config["JUDICEX_DEFAULT_MODEL"],
                ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
                overrides=payload,
            )
            try:
                result = test_provider(settings)
            except Exception as exc:
                return _error(str(exc), status=400)
            return jsonify({"result": result, "settings": public_settings(settings)})

    @app.post("/api/matters")
    def api_create_matter():
        payload = _json_payload()
        title = str(payload.get("title", "")).strip()
        if not title:
            return _error("Il titolo del fascicolo è obbligatorio.", status=400)
        with store() as memory:
            matter = memory.create_matter(
                title,
                client_name=str(payload.get("client_name", "")).strip(),
                area=str(payload.get("area", "")).strip(),
                status=str(payload.get("status", "open")).strip() or "open",
                summary=str(payload.get("summary", "")).strip(),
            )
            return jsonify({"matter": matter, "matters": memory.list_matters(top_k=100)})

    @app.get("/api/matters/<path:matter_id>")
    def api_get_matter(matter_id: str):
        with store() as memory:
            matter = memory.get_matter(matter_id)
            if matter is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            return jsonify({"matter": matter, "context": memory.build_matter_context(matter_id)})

    @app.patch("/api/matters/<path:matter_id>")
    def api_update_matter(matter_id: str):
        payload = _json_payload()
        with store() as memory:
            matter = memory.get_matter(matter_id)
            if matter is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            updated = memory.create_matter(
                title=str(payload.get("title", matter["title"])).strip() or matter["title"],
                client_name=str(payload.get("client_name", matter["client_name"])).strip(),
                area=str(payload.get("area", matter["area"])).strip(),
                status=str(payload.get("status", matter["status"])).strip() or "open",
                summary=str(payload.get("summary", matter["summary"])).strip(),
                metadata={**(matter.get("metadata") or {}), **(payload.get("metadata") if isinstance(payload.get("metadata"), dict) else {})},
                matter_id=matter_id,
            )
            return jsonify({"matter": updated, "matters": memory.list_matters(top_k=100)})

    @app.delete("/api/matters/<path:matter_id>")
    def api_delete_matter(matter_id: str):
        with store() as memory:
            if memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            deleted = memory.delete_matter(matter_id)
            if not deleted:
                return _error(f"Fascicolo non eliminato: {matter_id}", status=400)
            return jsonify({"status": "deleted", "matters": memory.list_matters(top_k=100)})

    @app.get("/api/matters/<path:matter_id>/documents")
    def api_list_matter_documents(matter_id: str):
        query = str(request.args.get("query", ""))
        with store() as memory:
            if memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            documents = memory.search_matter_documents(query, matter_id=matter_id, top_k=int(request.args.get("top_k", 50)))
            return jsonify({"documents": documents})

    @app.get("/api/matters/<path:matter_id>/export")
    def api_export_matter(matter_id: str):
        fmt = str(request.args.get("format", "docx")).lower().strip()
        with store() as memory:
            matter = memory.get_matter(matter_id)
            if matter is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            refs = memory.search_matter_documents("", matter_id=matter_id, top_k=100)
            documents = [
                doc
                for doc in (memory.get_matter_document(ref["id"], full=True) for ref in refs)
                if doc is not None
            ]
        if fmt == "zip":
            return _matter_zip_response(matter, documents)
        return _matter_docx_response(matter, documents)

    @app.post("/api/matters/<path:matter_id>/documents")
    def api_upload_matter_documents(matter_id: str):
        files = request.files.getlist("files")
        if not files:
            return _error("Nessun file ricevuto.", status=400)
        results: list[dict[str, Any]] = []
        with store() as memory:
            if memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            for upload in files:
                filename = Path(upload.filename or "documento.txt").name
                suffix = Path(filename).suffix.lower()
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    temp_path = Path(tmp.name)
                upload.save(temp_path)
                try:
                    result = memory.add_matter_document_file(
                        matter_id,
                        temp_path,
                        title=Path(filename).stem or filename,
                        kind=str(request.form.get("kind", "document") or "document"),
                        metadata={"uploaded_filename": filename},
                    )
                    results.append({"filename": filename, "status": "ok", **result})
                except Exception as exc:
                    results.append({"filename": filename, "status": "error", "error": str(exc)})
                finally:
                    try:
                        temp_path.unlink()
                    except OSError:
                        pass
            return jsonify(
                {
                    "uploads": results,
                    "context": memory.build_matter_context(matter_id),
                    "health": memory.health(),
                }
            )

    @app.get("/api/matters/<path:matter_id>/folders")
    def api_list_matter_folders(matter_id: str):
        with store() as memory:
            if memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            return jsonify({"folders": memory.list_matter_folders(matter_id)})

    @app.post("/api/matters/<path:matter_id>/folders")
    def api_create_matter_folder(matter_id: str):
        payload = _json_payload()
        with store() as memory:
            folder = memory.create_matter_folder(
                matter_id,
                str(payload.get("name", "")).strip(),
                parent_id=str(payload.get("parent_id", "")).strip(),
                sort_order=int(payload.get("sort_order", 0) or 0),
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            if "error" in folder:
                return _error(folder["error"], status=400 if "required" in folder["error"] else 404)
            return jsonify({"folder": folder, "folders": memory.list_matter_folders(matter_id)})

    @app.patch("/api/matter-folders/<path:folder_id>")
    def api_update_matter_folder(folder_id: str):
        payload = _json_payload()
        with store() as memory:
            folder = memory.update_matter_folder(
                folder_id,
                name=str(payload["name"]).strip() if "name" in payload else None,
                parent_id=str(payload["parent_id"]).strip() if "parent_id" in payload else None,
                sort_order=int(payload["sort_order"]) if "sort_order" in payload else None,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            if folder is None:
                return _error(f"Cartella non trovata: {folder_id}", status=404)
            return jsonify({"folder": folder})

    @app.get("/api/matter-documents/<path:document_id>")
    def api_get_matter_document(document_id: str):
        with store() as memory:
            document = memory.get_matter_document(document_id, full=True)
            if document is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            return jsonify(
                {
                    "document": document,
                    "versions": memory.list_matter_document_versions(document_id),
                    "edits": memory.list_document_edits(document_id),
                    "annotations": memory.list_document_annotations(document_id),
                    "comments": memory.list_document_comments(document_id),
                }
            )

    @app.get("/api/matter-documents/<path:document_id>/file")
    def api_matter_document_file(document_id: str):
        with store() as memory:
            document = memory.get_matter_document(document_id, full=True)
            if document is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            source_path = _safe_stored_path(document, Path(app.config["JUDICEX_DB_PATH"]).resolve().parent)
            if source_path is None:
                return _text_download_response(document["title"], document.get("content", ""), as_attachment=False)
            return send_file(source_path, as_attachment=False, download_name=str((document.get("metadata") or {}).get("original_filename") or source_path.name))

    @app.get("/api/matter-documents/<path:document_id>/preview")
    def api_matter_document_preview(document_id: str):
        with store() as memory:
            document = memory.get_matter_document(document_id, full=True)
            if document is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            source_path = _safe_stored_path(document, Path(app.config["JUDICEX_DB_PATH"]).resolve().parent)
            preview = _document_preview_payload(document, source_path)
            return jsonify(
                {
                    "document": document,
                    "preview": preview,
                    "annotations": memory.list_document_annotations(document_id),
                    "comments": memory.list_document_comments(document_id),
                    "versions": memory.list_matter_document_versions(document_id),
                    "edits": memory.list_document_edits(document_id),
                }
            )

    @app.post("/api/matter-documents/<path:document_id>/ocr")
    def api_matter_document_ocr(document_id: str):
        payload = _json_payload()
        with store() as memory:
            document = memory.get_matter_document(document_id, full=True)
            if document is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            source_path = _safe_stored_path(document, Path(app.config["JUDICEX_DB_PATH"]).resolve().parent)
            ocr = _extract_ocr_text(
                document,
                source_path,
                ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
                ocr_model=app.config["JUDICEX_OCR_MODEL"],
                max_pdf_pages=app.config["JUDICEX_OCR_MAX_PDF_PAGES"],
            )
            if ocr.get("text") and payload.get("apply", True):
                updated = memory.update_matter_document(
                    document_id,
                    content=str(ocr["text"]),
                    metadata={"ocr": {k: v for k, v in ocr.items() if k != "text"}},
                    reason="ocr",
                )
                return jsonify({"ocr": ocr, "document": updated})
            return jsonify({"ocr": ocr})

    @app.patch("/api/matter-documents/<path:document_id>")
    def api_update_matter_document(document_id: str):
        payload = _json_payload()
        with store() as memory:
            document = memory.update_matter_document(
                document_id,
                title=str(payload["title"]).strip() if "title" in payload else None,
                kind=str(payload["kind"]).strip() if "kind" in payload else None,
                content=str(payload["content"]) if "content" in payload else None,
                folder_id=str(payload["folder_id"]).strip() if "folder_id" in payload else None,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
                reason=str(payload.get("reason", "manual_update")),
            )
            if document is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            if "error" in document:
                return _error(document["error"], status=400)
            return jsonify({"document": document, "versions": memory.list_matter_document_versions(document_id)})

    @app.patch("/api/matter-documents/<path:document_id>/folder")
    def api_assign_matter_document_folder(document_id: str):
        payload = _json_payload()
        with store() as memory:
            document = memory.assign_matter_document_folder(document_id, str(payload.get("folder_id", "")).strip())
            if document is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            if "error" in document:
                return _error(document["error"], status=400)
            return jsonify({"document": document})

    @app.get("/api/matter-documents/<path:document_id>/versions")
    def api_list_matter_document_versions(document_id: str):
        with store() as memory:
            if memory.get_matter_document(document_id, full=False) is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            return jsonify({"versions": memory.list_matter_document_versions(document_id, full=bool(request.args.get("full")))})

    @app.post("/api/matter-documents/<path:document_id>/versions/<path:version_id>/restore")
    def api_restore_matter_document_version(document_id: str, version_id: str):
        with store() as memory:
            document = memory.restore_matter_document_version(document_id, version_id)
            if document is None:
                return _error(f"Versione non trovata: {version_id}", status=404)
            return jsonify({"document": document, "versions": memory.list_matter_document_versions(document_id)})

    @app.get("/api/matter-documents/<path:document_id>/versions/<path:version_id>/compare")
    def api_compare_matter_document_version(document_id: str, version_id: str):
        other = str(request.args.get("other", "")).strip()
        with store() as memory:
            comparison = memory.compare_matter_document_versions(document_id, version_id, other)
            if comparison is None:
                return _error(f"Versione non trovata: {version_id}", status=404)
            return jsonify({"comparison": comparison})

    @app.get("/api/matter-documents/<path:document_id>/download")
    def api_download_matter_document(document_id: str):
        file_format = str(request.args.get("format", "")).lower().strip()
        with store() as memory:
            document = memory.get_matter_document(document_id, full=True)
            if document is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            if file_format == "docx":
                return _docx_response(document["title"], document.get("content", ""))
            source_path = _safe_stored_path(document, Path(app.config["JUDICEX_DB_PATH"]).resolve().parent)
            if source_path is not None:
                download_name = str((document.get("metadata") or {}).get("original_filename") or source_path.name)
                return send_file(source_path, as_attachment=True, download_name=download_name)
            return _text_download_response(document["title"], document.get("content", ""))

    @app.post("/api/artifacts")
    def api_create_artifact():
        payload = _json_payload()
        title = str(payload.get("title", "")).strip() or "Documento Judicex"
        content = str(payload.get("content", "")).strip()
        if not content:
            return _error("Non c'e contenuto da trasformare in documento.", status=400)
        artifact_format = _normalize_artifact_format(payload.get("format"))
        session_id = str(payload.get("session_id", "")).strip()
        matter_id = str(payload.get("matter_id", "")).strip()
        with store() as memory:
            if matter_id and memory.get_matter(matter_id) is None:
                matter_id = ""
            artifact = memory.create_generated_artifact(
                title=title,
                content=content,
                format=artifact_format,
                session_id=session_id,
                matter_id=matter_id,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else {},
            )
            document = None
            if matter_id and bool(payload.get("save_to_matter")):
                stored = memory.add_matter_document(
                    matter_id,
                    title=title,
                    kind="draft",
                    content=content,
                    metadata={
                        "generated_by": "judicex_chat_artifact",
                        "artifact_id": artifact["id"],
                        "format": artifact_format,
                    },
                )
                document = stored.get("document") if isinstance(stored, dict) else None
            return jsonify({"artifact": artifact, "document": document})

    @app.get("/api/artifacts/<path:artifact_id>")
    def api_get_artifact(artifact_id: str):
        with store() as memory:
            artifact = memory.get_generated_artifact(artifact_id, full=True)
            if artifact is None:
                return _error(f"Documento generato non trovato: {artifact_id}", status=404)
            return jsonify({"artifact": artifact})

    @app.get("/api/artifacts/<path:artifact_id>/download")
    def api_download_artifact(artifact_id: str):
        requested_format = _normalize_artifact_format(request.args.get("format"))
        with store() as memory:
            artifact = memory.get_generated_artifact(artifact_id, full=True)
            if artifact is None:
                return _error(f"Documento generato non trovato: {artifact_id}", status=404)
            file_format = requested_format or _normalize_artifact_format(artifact.get("format")) or "docx"
            title = str(artifact.get("title") or "Documento Judicex")
            content = str(artifact.get("content") or "")
            if file_format == "docx":
                return _docx_response(title, content)
            if file_format == "pdf":
                return _pdf_response(title, content)
            if file_format == "md":
                return _text_download_response(title, content, extension="md", mimetype="text/markdown; charset=utf-8")
            return _text_download_response(title, content)

    @app.get("/api/matter-documents/<path:document_id>/edits")
    def api_list_document_edits(document_id: str):
        with store() as memory:
            if memory.get_matter_document(document_id, full=False) is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            return jsonify({"edits": memory.list_document_edits(document_id)})

    @app.get("/api/matter-documents/<path:document_id>/annotations")
    def api_list_document_annotations(document_id: str):
        with store() as memory:
            if memory.get_matter_document(document_id, full=False) is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            return jsonify({"annotations": memory.list_document_annotations(document_id)})

    @app.post("/api/matter-documents/<path:document_id>/annotations")
    def api_add_document_annotation(document_id: str):
        payload = _json_payload()
        with store() as memory:
            annotation = memory.add_document_annotation(
                document_id,
                page_number=int(payload.get("page_number", 1) or 1),
                x=float(payload.get("x", 0) or 0),
                y=float(payload.get("y", 0) or 0),
                width=float(payload.get("width", 0) or 0),
                height=float(payload.get("height", 0) or 0),
                color=str(payload.get("color", "#facc15")),
                note=str(payload.get("note", "")),
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            if annotation is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            return jsonify({"annotation": annotation, "annotations": memory.list_document_annotations(document_id)})

    @app.patch("/api/document-annotations/<path:annotation_id>")
    def api_update_document_annotation(annotation_id: str):
        payload = _json_payload()
        with store() as memory:
            annotation = memory.update_document_annotation(annotation_id, **payload)
            if annotation is None:
                return _error(f"Annotazione non trovata: {annotation_id}", status=404)
            return jsonify({"annotation": annotation})

    @app.delete("/api/document-annotations/<path:annotation_id>")
    def api_delete_document_annotation(annotation_id: str):
        with store() as memory:
            deleted = memory.delete_document_annotation(annotation_id)
            if not deleted:
                return _error(f"Annotazione non trovata: {annotation_id}", status=404)
            return jsonify({"deleted": True})

    @app.get("/api/matter-documents/<path:document_id>/comments")
    def api_list_document_comments(document_id: str):
        with store() as memory:
            if memory.get_matter_document(document_id, full=False) is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            return jsonify({"comments": memory.list_document_comments(document_id)})

    @app.post("/api/matter-documents/<path:document_id>/comments")
    def api_add_document_comment(document_id: str):
        payload = _json_payload()
        with store() as memory:
            comment = memory.add_document_comment(
                document_id,
                body=str(payload.get("body", "")),
                anchor=str(payload.get("anchor", "")),
                status=str(payload.get("status", "open")),
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            if comment is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            if "error" in comment:
                return _error(comment["error"], status=400)
            return jsonify({"comment": comment, "comments": memory.list_document_comments(document_id)})

    @app.patch("/api/document-comments/<path:comment_id>")
    def api_update_document_comment(comment_id: str):
        payload = _json_payload()
        with store() as memory:
            comment = memory.update_document_comment(
                comment_id,
                body=str(payload["body"]) if "body" in payload else None,
                status=str(payload["status"]) if "status" in payload else None,
            )
            if comment is None:
                return _error(f"Commento non trovato: {comment_id}", status=404)
            return jsonify({"comment": comment})

    @app.post("/api/matter-documents/<path:document_id>/edits")
    def api_create_document_edit(document_id: str):
        payload = _json_payload()
        revised_content = str(payload.get("revised_content", payload.get("content", "")))
        if not revised_content.strip():
            return _error("Inserisci il testo revisionato.", status=400)
        with store() as memory:
            edit = memory.create_document_edit(
                document_id,
                revised_content,
                title=str(payload.get("title", "")).strip(),
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            if edit is None:
                return _error(f"Documento non trovato: {document_id}", status=404)
            return jsonify({"edit": edit, "edits": memory.list_document_edits(document_id)})

    @app.get("/api/document-edits/<path:edit_id>")
    def api_get_document_edit(edit_id: str):
        with store() as memory:
            edit = memory.get_document_edit(edit_id, full=True)
            if edit is None:
                return _error(f"Revisione non trovata: {edit_id}", status=404)
            return jsonify({"edit": edit})

    @app.post("/api/document-edits/<path:edit_id>/apply")
    def api_apply_document_edit(edit_id: str):
        with store() as memory:
            result = memory.apply_document_edit(edit_id)
            if result is None:
                return _error(f"Revisione non trovata: {edit_id}", status=404)
            return jsonify(result)

    @app.get("/api/matters/<path:matter_id>/context")
    def api_matter_context(matter_id: str):
        query = str(request.args.get("query", ""))
        with store() as memory:
            context = memory.build_matter_context(matter_id, query=query)
            if "error" in context:
                return _error(context["error"], status=404)
            return jsonify({"context": context})

    @app.get("/api/workflows")
    def api_workflows():
        with store() as memory:
            return jsonify({"workflow_packs": memory.list_workflow_packs()})

    @app.post("/api/workflows")
    def api_create_workflow():
        payload = _json_payload()
        requirements = payload.get("requirements")
        if not isinstance(requirements, list):
            return _error("Aggiungi almeno un requisito al workflow.", status=400)
        with store() as memory:
            workflow = memory.create_custom_workflow_pack(
                label=str(payload.get("label", "")).strip(),
                profile_label=str(payload.get("profile_label", "")).strip(),
                pack_id=str(payload.get("id", "")).strip(),
                match_terms=_string_list(payload.get("match_terms", [])),
                requirements=requirements,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            if "error" in workflow:
                return _error(workflow["error"], status=400)
            return jsonify({"workflow": workflow, "workflow_packs": memory.list_workflow_packs()})

    @app.get("/api/workflows/<path:workflow_id>")
    def api_get_workflow(workflow_id: str):
        with store() as memory:
            workflow = memory.get_custom_workflow_pack(workflow_id)
            if workflow is None:
                return _error(f"Workflow personalizzato non trovato: {workflow_id}", status=404)
            return jsonify({"workflow": workflow, "versions": memory.list_custom_workflow_versions(workflow_id)})

    @app.patch("/api/workflows/<path:workflow_id>")
    def api_update_workflow(workflow_id: str):
        payload = _json_payload()
        requirements = payload.get("requirements") if "requirements" in payload else None
        if requirements is not None and not isinstance(requirements, list):
            return _error("I requisiti devono essere una lista.", status=400)
        with store() as memory:
            workflow = memory.update_custom_workflow_pack(
                workflow_id,
                label=str(payload["label"]).strip() if "label" in payload else None,
                profile_label=str(payload.get("profile_label", "")).strip(),
                match_terms=_string_list(payload["match_terms"]) if "match_terms" in payload else None,
                requirements=requirements,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
                reason=str(payload.get("reason", "manual_update")),
            )
            if workflow is None:
                return _error(f"Workflow personalizzato non trovato: {workflow_id}", status=404)
            if "error" in workflow:
                return _error(workflow["error"], status=400)
            return jsonify({"workflow": workflow, "versions": memory.list_custom_workflow_versions(workflow_id), "workflow_packs": memory.list_workflow_packs()})

    @app.post("/api/workflows/<path:workflow_id>/duplicate")
    def api_duplicate_workflow(workflow_id: str):
        payload = _json_payload()
        with store() as memory:
            workflow = memory.duplicate_custom_workflow_pack(workflow_id, label=str(payload.get("label", "")).strip())
            if workflow is None:
                return _error(f"Workflow personalizzato non trovato: {workflow_id}", status=404)
            if "error" in workflow:
                return _error(workflow["error"], status=400)
            return jsonify({"workflow": workflow, "workflow_packs": memory.list_workflow_packs()})

    @app.delete("/api/workflows/<path:workflow_id>")
    def api_delete_workflow(workflow_id: str):
        with store() as memory:
            deleted = memory.delete_custom_workflow_pack(workflow_id)
            if not deleted:
                return _error(f"Workflow personalizzato non trovato: {workflow_id}", status=404)
            return jsonify({"deleted": True, "workflow_packs": memory.list_workflow_packs()})

    @app.post("/api/workflows/run")
    def api_run_workflow():
        payload = _json_payload()
        matter_id = str(payload.get("matter_id", "")).strip()
        thesis = str(payload.get("thesis", "")).strip() or "analisi generale del fascicolo"
        if not matter_id:
            return _error("Seleziona un fascicolo.", status=400)
        with store() as memory:
            result = memory.analyze_matter(
                matter_id,
                thesis,
                workflow_pack=str(payload.get("workflow_pack", "")).strip() or None,
            )
            return jsonify({"analysis": result})

    @app.get("/api/matters/<path:matter_id>/tabular-reviews")
    def api_list_tabular_reviews(matter_id: str):
        with store() as memory:
            if memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            return jsonify({"reviews": memory.list_tabular_reviews(matter_id)})

    @app.get("/api/matters/<path:matter_id>/visualizations")
    def api_matter_visualizations(matter_id: str):
        with store() as memory:
            payload = memory.build_matter_visualizations(matter_id)
            if "error" in payload:
                return _error(payload["error"], status=404)
            return jsonify({"visualizations": payload})

    @app.post("/api/matters/<path:matter_id>/tabular-reviews")
    def api_create_tabular_review(matter_id: str):
        payload = _json_payload()
        with store() as memory:
            review = memory.create_tabular_review(
                matter_id,
                title=str(payload.get("title", "")).strip(),
                query=str(payload.get("query", "")).strip(),
            )
            if "error" in review:
                return _error(review["error"], status=404)
            return jsonify({"review": review, "reviews": memory.list_tabular_reviews(matter_id)})

    @app.get("/api/tabular-reviews/<path:review_id>")
    def api_get_tabular_review(review_id: str):
        with store() as memory:
            review = memory.get_tabular_review(review_id, full=True)
            if review is None:
                return _error(f"Revisione tabellare non trovata: {review_id}", status=404)
            return jsonify({"review": review, "views": memory.list_tabular_review_views(review_id)})

    @app.patch("/api/tabular-reviews/<path:review_id>")
    def api_update_tabular_review(review_id: str):
        payload = _json_payload()
        with store() as memory:
            review = memory.update_tabular_review(
                review_id,
                title=str(payload["title"]).strip() if "title" in payload else None,
                columns=payload.get("columns") if isinstance(payload.get("columns"), list) else None,
                rows=payload.get("rows") if isinstance(payload.get("rows"), list) else None,
            )
            if review is None:
                return _error(f"Revisione tabellare non trovata: {review_id}", status=404)
            return jsonify({"review": review, "views": memory.list_tabular_review_views(review_id)})

    @app.patch("/api/tabular-reviews/<path:review_id>/cell")
    def api_update_tabular_review_cell(review_id: str):
        payload = _json_payload()
        with store() as memory:
            review = memory.update_tabular_review_cell(
                review_id,
                int(payload.get("row_index", -1)),
                str(payload.get("key", "")).strip(),
                payload.get("value", ""),
            )
            if review is None:
                return _error(f"Revisione tabellare non trovata: {review_id}", status=404)
            if "error" in review:
                return _error(review["error"], status=400)
            return jsonify({"review": review})

    @app.post("/api/tabular-reviews/<path:review_id>/views")
    def api_save_tabular_review_view(review_id: str):
        payload = _json_payload()
        with store() as memory:
            view = memory.save_tabular_review_view(
                review_id,
                name=str(payload.get("name", "")).strip(),
                filter_text=str(payload.get("filter_text", "")).strip(),
                sort_key=str(payload.get("sort_key", "")).strip(),
                sort_dir=str(payload.get("sort_dir", "asc")).strip(),
                columns=payload.get("columns") if isinstance(payload.get("columns"), list) else None,
            )
            if view is None:
                return _error(f"Revisione tabellare non trovata: {review_id}", status=404)
            return jsonify({"view": view, "views": memory.list_tabular_review_views(review_id)})

    @app.get("/api/tabular-reviews/<path:review_id>/export")
    def api_export_tabular_review(review_id: str):
        fmt = str(request.args.get("format", "csv")).lower().strip()
        with store() as memory:
            review = memory.get_tabular_review(review_id, full=True)
            if review is None:
                return _error(f"Revisione tabellare non trovata: {review_id}", status=404)
        if fmt == "docx":
            return _review_docx_response(review)
        if fmt == "xlsx":
            return _review_xlsx_response(review)
        return _review_csv_response(review)

    @app.get("/api/draft-templates")
    def api_draft_templates():
        with store() as memory:
            custom = memory.list_custom_draft_templates()
        builtin = [{**item, "source": "builtin"} for item in list_templates()]
        return jsonify({"templates": builtin + custom})

    @app.post("/api/draft-templates")
    def api_save_draft_template():
        payload = _json_payload()
        body = str(payload.get("body", ""))
        title = str(payload.get("title", "")).strip()
        if not title or not body.strip():
            return _error("Titolo e corpo template sono obbligatori.", status=400)
        required_params = _string_list(payload.get("required_params", []))
        with store() as memory:
            template = memory.save_custom_draft_template(
                name=str(payload.get("name", title)),
                title=title,
                body=body,
                required_params=required_params,
                metadata=payload.get("metadata") if isinstance(payload.get("metadata"), dict) else None,
            )
            return jsonify({"template": template, "templates": [{**item, "source": "builtin"} for item in list_templates()] + memory.list_custom_draft_templates()})

    @app.get("/api/draft-templates/<path:template_id>")
    def api_get_draft_template(template_id: str):
        with store() as memory:
            template = memory.get_custom_draft_template(template_id)
            if template is not None:
                return jsonify({"template": template})
        return _error(f"Template personalizzato non trovato: {template_id}", status=404)

    @app.post("/api/draft-templates/<path:template_id>/preview")
    def api_preview_draft_template(template_id: str):
        payload = _json_payload()
        params = payload.get("params") if isinstance(payload.get("params"), dict) else {}
        with store() as memory:
            template = memory.get_custom_draft_template(template_id)
            if template is None:
                return _error(f"Template personalizzato non trovato: {template_id}", status=404)
            try:
                rendered = _render_custom_template(template, {str(k): str(v) for k, v in params.items()})
            except DraftingError as exc:
                return _error(str(exc), status=400)
            return jsonify({"template": template, "rendered": rendered})

    @app.post("/api/matters/<path:matter_id>/drafts/assistant")
    def api_assistant_draft(matter_id: str):
        payload = _json_payload()
        instruction = str(payload.get("instruction", "")).strip()
        template_name = str(payload.get("template_name", "")).strip()
        as_of_date = str(payload.get("as_of_date", "")).strip() or datetime.now(timezone.utc).date().isoformat()
        if not instruction:
            return _error("Scrivi cosa deve preparare Judicex.", status=400)
        with store() as memory:
            matter = memory.get_matter(matter_id)
            if matter is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            custom_templates = memory.list_custom_draft_templates()
            if not template_name:
                template_name = _choose_draft_template(instruction, custom_templates)
            required = _draft_required_params(memory, template_name)
            context = memory.build_matter_context(matter_id)
            params = _infer_draft_params(instruction, required, matter=matter, context=context)
            missing = [key for key in required if not params.get(key)]
            if missing:
                return jsonify(
                    {
                        "draft": {
                            "status": "needs_info",
                            "template": template_name,
                            "title": _draft_template_title(memory, template_name),
                            "reason": "Mi servono alcuni dati prima di preparare l'atto.",
                            "missing_params": missing,
                            "suggested_params": params,
                        },
                        "context": context,
                    }
                )

            custom_template = memory.get_custom_draft_template(template_name)
            if custom_template is not None:
                rendered = _render_custom_template(custom_template, params)
                stored = memory.add_matter_document(
                    matter_id,
                    title=str(payload.get("title", "")).strip() or custom_template["title"],
                    kind="draft",
                    content=rendered,
                    metadata={
                        "generated_by": "judicex_assistant_custom_template",
                        "template_name": custom_template["name"],
                        "as_of_date": as_of_date,
                        "instruction": instruction,
                    },
                )
                return jsonify(
                    {
                        "draft": {
                            "status": "drafted",
                            "template": custom_template["name"],
                            "title": custom_template["title"],
                            "as_of_date": as_of_date,
                            "rendered": rendered,
                            "citations": [],
                            "params": params,
                        },
                        "document": stored.get("document"),
                        "context": memory.build_matter_context(matter_id),
                    }
                )

            try:
                draft = draft_atto(
                    memory,
                    template_name=template_name,
                    as_of_date=as_of_date,
                    params=params,
                    matter_id=matter_id,
                )
            except DraftingError as exc:
                return _error(_friendly_draft_error(str(exc)), status=400)
            if draft.get("status") != "drafted":
                return jsonify({"draft": draft, "context": context})
            stored = memory.add_matter_document(
                matter_id,
                title=str(payload.get("title", "")).strip() or str(draft.get("title") or template_name),
                kind="draft",
                content=str(draft.get("rendered", "")),
                metadata={
                    "generated_by": "judicex_assistant_drafter",
                    "template_name": template_name,
                    "as_of_date": as_of_date,
                    "instruction": instruction,
                    "citations": draft.get("citations", []),
                    "warnings": draft.get("warnings", []),
                },
            )
            draft["params"] = params
            return jsonify({"draft": draft, "document": stored.get("document"), "context": memory.build_matter_context(matter_id)})

    @app.post("/api/matters/<path:matter_id>/drafts")
    def api_create_draft(matter_id: str):
        payload = _json_payload()
        template_name = str(payload.get("template_name", "")).strip()
        if not template_name:
            return _error("Seleziona un template atto.", status=400)
        params = payload.get("params")
        if isinstance(params, str):
            try:
                params = _json_loads_object(params)
            except ValueError as exc:
                return _error(str(exc), status=400)
        if params is not None and not isinstance(params, dict):
            return _error("I dati del documento devono essere scritti come nome: valore.", status=400)
        as_of_date = str(payload.get("as_of_date", "")).strip() or datetime.now(timezone.utc).date().isoformat()
        with store() as memory:
            if memory.get_matter(matter_id) is None:
                return _error(f"Fascicolo non trovato: {matter_id}", status=404)
            custom_template = memory.get_custom_draft_template(template_name)
            if custom_template is not None:
                params_dict = {str(k): str(v) for k, v in (params or {}).items()}
                missing = [key for key in custom_template.get("required_params", []) if not params_dict.get(key)]
                if missing:
                    return _error(f"Parametri mancanti: {', '.join(missing)}", status=400)
                rendered = _render_custom_template(custom_template, params_dict)
                stored = memory.add_matter_document(
                    matter_id,
                    title=str(payload.get("title", "")).strip() or custom_template["title"],
                    kind="draft",
                    content=rendered,
                    metadata={
                        "generated_by": "judicex_custom_template",
                        "template_name": custom_template["name"],
                        "as_of_date": as_of_date,
                    },
                )
                return jsonify(
                    {
                        "draft": {
                            "status": "drafted",
                            "template": custom_template["name"],
                            "title": custom_template["title"],
                            "as_of_date": as_of_date,
                            "rendered": rendered,
                            "citations": [],
                        },
                        "document": stored.get("document"),
                        "context": memory.build_matter_context(matter_id),
                    }
                )
            try:
                draft = draft_atto(
                    memory,
                    template_name=template_name,
                    as_of_date=as_of_date,
                    params={str(k): str(v) for k, v in (params or {}).items()},
                    matter_id=matter_id,
                )
            except DraftingError as exc:
                return _error(str(exc), status=400)
            if draft.get("status") != "drafted":
                return jsonify({"draft": draft})
            stored = memory.add_matter_document(
                matter_id,
                title=str(payload.get("title", "")).strip() or str(draft.get("title") or template_name),
                kind="draft",
                content=str(draft.get("rendered", "")),
                metadata={
                    "generated_by": "judicex_drafter",
                    "template_name": template_name,
                    "as_of_date": as_of_date,
                    "citations": draft.get("citations", []),
                    "warnings": draft.get("warnings", []),
                },
            )
            return jsonify({"draft": draft, "document": stored.get("document"), "context": memory.build_matter_context(matter_id)})

    @app.get("/api/tools")
    def api_tools():
        with store() as memory:
            return jsonify({"tools": LegalMemoryTools(memory).definitions()})

    @app.post("/api/tools/call")
    def api_call_tool():
        payload = _json_payload()
        name = str(payload.get("name", "")).strip()
        arguments = payload.get("arguments") if isinstance(payload.get("arguments"), dict) else {}
        if not name:
            return _error("Nome tool obbligatorio.", status=400)
        with store() as memory:
            try:
                result = LegalMemoryTools(memory).call(name, arguments)
            except Exception as exc:
                return _error(str(exc), status=400)
            return jsonify({"result": result})

    @app.post("/api/analyze")
    def api_analyze():
        payload = _json_payload()
        matter_id = str(payload.get("matter_id", "")).strip()
        thesis = str(payload.get("thesis", "")).strip()
        if not matter_id:
            return _error("Seleziona un fascicolo.", status=400)
        if not thesis:
            return _error("Inserisci una tesi o obiettivo di analisi.", status=400)
        with store() as memory:
            result = memory.analyze_matter(
                matter_id,
                thesis,
                workflow_pack=str(payload.get("workflow_pack", "")).strip() or None,
            )
            return jsonify({"analysis": result})

    @app.post("/api/chat")
    def api_chat():
        payload = _json_payload()
        question = str(payload.get("question", "")).strip()
        if not question:
            return _error("Scrivi una domanda.", status=400)
        with store() as memory:
            settings = resolve_settings(
                memory,
                default_model=app.config["JUDICEX_DEFAULT_MODEL"],
                ollama_host=app.config["JUDICEX_OLLAMA_HOST"],
                overrides=payload.get("llm") if isinstance(payload.get("llm"), dict) else None,
            )
            if not settings["model"] and settings["provider"] != "none":
                return _error("Configura provider e modello AI in Impostazioni.", status=400)
            client = make_client(settings)
            engine = JudicexAgentRuntime(
                store=memory,
                client=client,
                model=settings["model"],
                area=str(payload.get("area", "") or app.config["JUDICEX_DEFAULT_AREA"]).strip() or None,
                matter_id=str(payload.get("matter_id", "")).strip() or None,
            )
            result = engine.answer(
                question,
                recent_user_turns=[
                    str(item).strip()
                    for item in payload.get("recent_user_turns", [])
                    if str(item).strip()
                ],
            )
            return jsonify({"result": result})

    @app.post("/api/answer/stream")
    def api_answer_stream():
        """Server-Sent Events: streams agent_trace steps + final payload.

        The engine runs in a background thread; the main request thread polls
        a shared list and emits each new step as a `step` event. When the
        engine returns, a final `result` event is sent and the stream ends.
        """

        from .stream_runner import stream_answer

        payload = _json_payload()
        question = str(payload.get("question", "")).strip()
        if not question:
            return _error("Scrivi una domanda.", status=400)
        area = str(payload.get("area", "") or app.config["JUDICEX_DEFAULT_AREA"]).strip() or None
        matter_id = str(payload.get("matter_id", "")).strip() or None
        recent_user_turns = [
            str(item).strip()
            for item in payload.get("recent_user_turns", [])
            if str(item).strip()
        ]
        llm_payload = payload.get("llm") if isinstance(payload.get("llm"), dict) else {}

        def event_stream():
            for chunk in stream_answer(
                db_path=app.config["JUDICEX_DB_PATH"],
                question=question,
                model=str(llm_payload.get("model") or app.config["JUDICEX_DEFAULT_MODEL"]).strip(),
                host=str(llm_payload.get("base_url") or app.config["JUDICEX_OLLAMA_HOST"]).strip(),
                provider=str(llm_payload.get("provider") or "").strip(),
                base_url=str(llm_payload.get("base_url") or "").strip(),
                area=area,
                matter_id=matter_id,
                recent_user_turns=recent_user_turns,
            ):
                yield chunk

        return Response(event_stream(), mimetype="text/event-stream", headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        })

    return app


def _json_payload() -> dict[str, Any]:
    payload = request.get_json(silent=True)
    return payload if isinstance(payload, dict) else {}


def _error(message: str, *, status: int) -> tuple[Any, int]:
    return jsonify({"error": message}), status


def _json_loads_object(raw: str) -> dict[str, Any]:
    try:
        payload = json.loads(raw or "{}")
    except json.JSONDecodeError as exc:
        raise ValueError("Formato dati non valido. Scrivi nome: valore, uno per riga.") from exc
    if not isinstance(payload, dict):
        raise ValueError("Formato dati non valido. Scrivi nome: valore, uno per riga.")
    return payload


def _string_list(raw: Any) -> list[str]:
    if isinstance(raw, str):
        parts = re.split(r"[,;\n]+", raw)
    elif isinstance(raw, list):
        parts = raw
    else:
        parts = []
    return [str(part).strip() for part in parts if str(part).strip()]


def _safe_filename(title: str, extension: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", title.strip()).strip("._") or "judicex_document"
    return f"{stem[:80]}.{extension.lstrip('.')}"


def _normalize_artifact_format(value: Any) -> str:
    raw = re.sub(r"[^a-z0-9]+", "", str(value or "").strip().lower())
    aliases = {
        "doc": "docx",
        "word": "docx",
        "document": "docx",
        "documento": "docx",
        "markdown": "md",
        "text": "txt",
        "testo": "txt",
    }
    normalized = aliases.get(raw, raw)
    return normalized if normalized in {"docx", "pdf", "txt", "md"} else ""


def _text_download_response(
    title: str,
    content: str,
    *,
    as_attachment: bool = True,
    extension: str = "txt",
    mimetype: str = "text/plain; charset=utf-8",
):
    data = io.BytesIO(str(content or "").encode("utf-8"))
    return send_file(
        data,
        mimetype=mimetype,
        as_attachment=as_attachment,
        download_name=_safe_filename(title, extension),
    )


def _docx_response(title: str, content: str):
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx non è installato.") from exc
    doc = DocxDocument()
    clean_title = title.strip() or "Judicex"
    doc.add_heading(clean_title, level=1)
    doc.add_paragraph(f"Esportato da Judicex il {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    _add_docx_body(doc, str(content or ""))
    data = io.BytesIO()
    doc.save(data)
    data.seek(0)
    return send_file(
        data,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=_safe_filename(clean_title, "docx"),
    )


def _add_docx_body(doc: Any, content: str) -> None:
    for raw_block in re.split(r"\n{2,}", str(content or "").strip()):
        block = raw_block.strip()
        if not block:
            continue
        lines = [line.rstrip() for line in block.splitlines()]
        if len(lines) == 1:
            text = lines[0].strip()
            if text.isupper() and len(text) <= 140:
                doc.add_heading(text.title(), level=2)
            elif text.endswith(":") and len(text) <= 100:
                doc.add_heading(text[:-1], level=3)
            elif re.match(r"^[-*]\s+", text):
                doc.add_paragraph(re.sub(r"^[-*]\s+", "", text), style="List Bullet")
            else:
                doc.add_paragraph(text)
            continue
        for line in lines:
            text = line.strip()
            if not text:
                continue
            if re.match(r"^[-*]\s+", text):
                doc.add_paragraph(re.sub(r"^[-*]\s+", "", text), style="List Bullet")
            else:
                doc.add_paragraph(text)


def _pdf_response(title: str, content: str):
    data = io.BytesIO(_build_basic_pdf(title, content))
    return send_file(
        data,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=_safe_filename(title, "pdf"),
    )


def _build_basic_pdf(title: str, content: str) -> bytes:
    page_width = 612
    page_height = 792
    margin = 54
    leading = 14
    lines_per_page = max(1, int((page_height - (margin * 2)) / leading))
    lines = _pdf_wrapped_lines(str(title or "Documento Judicex"), width=80)
    lines.extend([""])
    lines.extend(_pdf_wrapped_lines(str(content or ""), width=92))
    pages = [lines[index:index + lines_per_page] for index in range(0, len(lines), lines_per_page)] or [[""]]

    objects: dict[int, bytes] = {}
    page_object_ids: list[int] = []
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[3] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    next_object = 4
    for page_lines in pages:
        page_object_id = next_object
        content_object_id = next_object + 1
        next_object += 2
        page_object_ids.append(page_object_id)
        stream = _pdf_content_stream(page_lines, margin, page_height - margin, leading)
        objects[content_object_id] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        )
        objects[page_object_id] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {page_width} {page_height}] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_object_id} 0 R >>"
        ).encode("ascii")
    kids = " ".join(f"{page_id} 0 R" for page_id in page_object_ids)
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_object_ids)} >>".encode("ascii")

    total_objects = max(objects)
    output = io.BytesIO()
    output.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets: list[int] = []
    for object_id in range(1, total_objects + 1):
        offsets.append(output.tell())
        output.write(f"{object_id} 0 obj\n".encode("ascii"))
        output.write(objects[object_id])
        output.write(b"\nendobj\n")
    xref_offset = output.tell()
    output.write(f"xref\n0 {total_objects + 1}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for offset in offsets:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.write(
        f"trailer\n<< /Size {total_objects + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return output.getvalue()


def _pdf_wrapped_lines(text: str, *, width: int) -> list[str]:
    wrapped: list[str] = []
    for raw_line in str(text or "").replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            wrapped.append("")
            continue
        words = line.split()
        current = ""
        for word in words:
            if len(word) > width:
                if current:
                    wrapped.append(current)
                    current = ""
                for index in range(0, len(word), width):
                    wrapped.append(word[index:index + width])
                continue
            candidate = word if not current else f"{current} {word}"
            if len(candidate) <= width:
                current = candidate
            else:
                wrapped.append(current)
                current = word
        if current:
            wrapped.append(current)
    return wrapped


def _pdf_content_stream(lines: list[str], x: int, y: int, leading: int) -> bytes:
    parts = ["BT", "/F1 10 Tf", f"{x} {y} Td", f"{leading} TL"]
    for line in lines:
        parts.append(f"({_pdf_escape_text(line)}) Tj")
        parts.append("T*")
    parts.append("ET")
    normalized = "\n".join(parts)
    return _pdf_latin_text(normalized).encode("latin-1", errors="replace")


def _pdf_escape_text(text: str) -> str:
    value = _pdf_latin_text(text)
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _pdf_latin_text(text: str) -> str:
    return (
        str(text or "")
        .replace("€", "EUR")
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
    )


def _matter_docx_response(matter: dict[str, Any], documents: list[dict[str, Any]]):
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx non è installato.") from exc
    doc = DocxDocument()
    title = matter.get("title") or "Fascicolo Judicex"
    doc.add_heading(title, level=1)
    meta = doc.add_paragraph()
    meta.add_run("Cliente: ").bold = True
    meta.add_run(str(matter.get("client_name") or ""))
    meta.add_run("\nArea: ").bold = True
    meta.add_run(str(matter.get("area") or ""))
    meta.add_run("\nEsportato: ").bold = True
    meta.add_run(datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))
    if matter.get("summary"):
        doc.add_heading("Sintesi", level=2)
        doc.add_paragraph(str(matter["summary"]))
    doc.add_heading("Documenti", level=2)
    for index, item in enumerate(documents, start=1):
        doc.add_heading(f"{index}. {item.get('title') or 'Documento'}", level=3)
        doc.add_paragraph(f"Tipo: {item.get('kind') or 'documento'}")
        _add_docx_body(doc, str(item.get("content") or ""))
    data = io.BytesIO()
    doc.save(data)
    data.seek(0)
    return send_file(
        data,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=_safe_filename(str(title), "docx"),
    )


def _matter_zip_response(matter: dict[str, Any], documents: list[dict[str, Any]]):
    data = io.BytesIO()
    with zipfile.ZipFile(data, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "matter.json",
            json.dumps({"matter": matter, "documents": [_document_export_meta(item) for item in documents]}, ensure_ascii=False, indent=2),
        )
        for index, item in enumerate(documents, start=1):
            stem = _safe_filename(f"{index:03d}_{item.get('title') or 'documento'}", "txt")
            archive.writestr(f"documents/{stem}", str(item.get("content") or ""))
    data.seek(0)
    return send_file(
        data,
        mimetype="application/zip",
        as_attachment=True,
        download_name=_safe_filename(str(matter.get("title") or "fascicolo"), "zip"),
    )


def _document_export_meta(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": item.get("id"),
        "title": item.get("title"),
        "kind": item.get("kind"),
        "created_at": item.get("created_at"),
        "updated_at": item.get("updated_at"),
        "metadata": item.get("metadata") or {},
    }


def _safe_stored_path(document: dict[str, Any], db_parent: Path) -> Path | None:
    metadata = document.get("metadata") or {}
    raw_path = str(metadata.get("stored_path") or document.get("source_path") or "").strip()
    if not raw_path:
        return None
    try:
        path = Path(raw_path).expanduser().resolve()
        base = db_parent.resolve()
    except OSError:
        return None
    if not path.is_file():
        return None
    if path == base or base in path.parents:
        return path
    return None


def _document_preview_payload(document: dict[str, Any], source_path: Path | None) -> dict[str, Any]:
    metadata = document.get("metadata") or {}
    suffix = source_path.suffix.lower() if source_path else Path(str(metadata.get("original_filename", ""))).suffix.lower()
    page_count = 1
    text_by_page: list[dict[str, Any]] = []
    viewer = "text"
    if source_path and suffix == ".pdf":
        viewer = "pdf"
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(source_path))
            page_count = len(reader.pages)
            for index, page in enumerate(reader.pages[:50], start=1):
                text_by_page.append({"page": index, "text": (page.extract_text() or "").strip()})
        except Exception as exc:
            text_by_page.append({"page": 1, "text": f"Errore lettura PDF: {exc}"})
    elif source_path and suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp", ".gif"}:
        viewer = "image"
    return {
        "viewer": viewer,
        "source_url": f"/api/matter-documents/{document['id']}/file" if source_path else "",
        "suffix": suffix,
        "page_count": page_count,
        "text_by_page": text_by_page,
        "content": document.get("content", ""),
        "metadata": metadata,
    }


def _extract_ocr_text(
    document: dict[str, Any],
    source_path: Path | None,
    *,
    ollama_host: str,
    ocr_model: str = "glm-5:cloud",
    max_pdf_pages: int = 20,
) -> dict[str, Any]:
    if source_path is None:
        return {"status": "no_source_file", "text": document.get("content", ""), "engine": "stored_text"}
    suffix = source_path.suffix.lower()
    if suffix == ".pdf":
        text_chunks: list[str] = []
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(source_path))
            for page in reader.pages:
                text_chunks.append(page.extract_text() or "")
            text = "\n\n".join(chunk.strip() for chunk in text_chunks if chunk.strip())
            if text.strip():
                return {"status": "extracted", "text": text, "engine": "pypdf", "pages": len(reader.pages)}
            return _ocr_pdf_with_ollama(source_path, ollama_host=ollama_host, model=ocr_model, max_pages=max_pdf_pages, page_count=len(reader.pages))
        except Exception as exc:
            return {"status": "error", "text": "", "engine": "pypdf", "error": str(exc)}
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
        ollama_result = _ocr_image_file_with_ollama(source_path, ollama_host=ollama_host, model=ocr_model)
        if ollama_result.get("text"):
            return ollama_result
        tesseract_result = _ocr_image_with_tesseract(source_path)
        if tesseract_result.get("text"):
            tesseract_result["fallback_from"] = {k: v for k, v in ollama_result.items() if k != "text"}
            return tesseract_result
        if ollama_result.get("status") != "ocr_done":
            return ollama_result
        return tesseract_result
    return {"status": "unsupported_format", "text": document.get("content", ""), "engine": "stored_text"}


def _ocr_pdf_with_ollama(source_path: Path, *, ollama_host: str, model: str, max_pages: int, page_count: int) -> dict[str, Any]:
    try:
        import fitz  # type: ignore
    except ModuleNotFoundError:
        return {
            "status": "scanned_pdf_no_text",
            "text": "",
            "engine": "ollama",
            "model": model,
            "pages": page_count,
            "note": "PDF senza testo incorporato. Per leggerlo con glm-5:cloud installa l'extra OCR, poi riprova.",
        }
    try:
        pdf = fitz.open(str(source_path))
    except Exception as exc:
        return {"status": "error", "text": "", "engine": "ollama", "model": model, "error": str(exc)}

    page_texts: list[str] = []
    errors: list[str] = []
    total_pages = len(pdf)
    if total_pages < 1:
        pdf.close()
        return {"status": "ocr_empty", "text": "", "engine": "ollama", "model": model, "pages": 0, "note": "PDF senza pagine leggibili."}
    limit = max(1, min(_safe_positive_int(max_pages, default=20), total_pages))
    try:
        for index in range(limit):
            page = pdf.load_page(index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            result = _ocr_image_bytes_with_ollama(
                pixmap.tobytes("png"),
                ollama_host=ollama_host,
                model=model,
                label=f"pagina {index + 1}",
            )
            if result.get("text"):
                page_texts.append(f"[Pagina {index + 1}]\n{result['text']}")
            elif result.get("error") or result.get("note"):
                errors.append(str(result.get("error") or result.get("note")))
    finally:
        pdf.close()

    text = "\n\n".join(page_texts).strip()
    if text:
        payload: dict[str, Any] = {"status": "ocr_done", "text": text, "engine": "ollama", "model": model, "pages": page_count, "processed_pages": limit}
        if limit < page_count:
            payload["note"] = f"OCR applicato alle prime {limit} pagine su {page_count}. Aumenta JUDICEX_OCR_MAX_PDF_PAGES per elaborarne di più."
        return payload
    return {
        "status": "ocr_empty",
        "text": "",
        "engine": "ollama",
        "model": model,
        "pages": page_count,
        "processed_pages": limit,
        "note": errors[0] if errors else "glm-5:cloud non ha restituito testo leggibile.",
    }


def _ocr_image_file_with_ollama(source_path: Path, *, ollama_host: str, model: str) -> dict[str, Any]:
    try:
        return _ocr_image_bytes_with_ollama(source_path.read_bytes(), ollama_host=ollama_host, model=model, label=source_path.name)
    except OSError as exc:
        return {"status": "error", "text": "", "engine": "ollama", "model": model, "error": str(exc)}


def _ocr_image_bytes_with_ollama(image_bytes: bytes, *, ollama_host: str, model: str, label: str = "immagine") -> dict[str, Any]:
    prompt = (
        "Trascrivi fedelmente tutto il testo visibile in questa immagine. "
        "Mantieni righe, numeri, date, importi, nomi e punteggiatura. "
        "Non riassumere, non commentare, non aggiungere spiegazioni. "
        "Se una parola non e leggibile usa [illeggibile]. Restituisci solo la trascrizione."
    )
    image_b64 = base64.b64encode(image_bytes).decode("ascii")
    try:
        from .ollama_agent import OllamaClient

        client = OllamaClient(host=ollama_host, timeout=600)
        text = client.chat(
            model=model,
            messages=[{"role": "user", "content": prompt, "images": [image_b64]}],
            temperature=0,
        )
    except Exception as exc:
        return {
            "status": "ocr_engine_missing",
            "text": "",
            "engine": "ollama",
            "model": model,
            "error": str(exc),
            "note": f"Avvia Ollama e rendi disponibile il modello {model} per usare OCR AI.",
        }
    clean_text = _clean_ocr_output(text)
    return {"status": "ocr_done", "text": clean_text, "engine": "ollama", "model": model, "source": label}


def _ocr_image_with_tesseract(source_path: Path) -> dict[str, Any]:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return {
            "status": "ocr_engine_missing",
            "text": "",
            "engine": "tesseract",
            "note": "Tesseract non è installato nel sistema.",
        }
    try:
        proc = subprocess.run(
            [tesseract, str(source_path), "stdout", "-l", "ita+eng"],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
    except Exception as exc:
        return {"status": "error", "text": "", "engine": "tesseract", "error": str(exc)}
    if proc.returncode != 0:
        return {"status": "error", "text": "", "engine": "tesseract", "error": proc.stderr.strip()}
    return {"status": "ocr_done", "text": proc.stdout.strip(), "engine": "tesseract"}


def _clean_ocr_output(text: str) -> str:
    value = str(text or "").strip()
    if value.startswith("```"):
        value = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", value)
        value = re.sub(r"\s*```$", "", value)
    return value.strip()


def _review_csv_response(review: dict[str, Any]):
    output = io.StringIO()
    writer = csv.writer(output)
    columns = review.get("columns") or []
    rows = review.get("rows") or []
    writer.writerow([column.get("label") or column.get("key") for column in columns])
    for row in rows:
        writer.writerow([row.get(column.get("key"), "") for column in columns])
    data = io.BytesIO(output.getvalue().encode("utf-8-sig"))
    return send_file(
        data,
        mimetype="text/csv; charset=utf-8",
        as_attachment=True,
        download_name=_safe_filename(review.get("title", "tabella"), "csv"),
    )


def _review_docx_response(review: dict[str, Any]):
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("python-docx non è installato.") from exc
    columns = review.get("columns") or []
    rows = review.get("rows") or []
    doc = DocxDocument()
    doc.add_heading(str(review.get("title") or "Tabella Judicex"), level=1)
    table = doc.add_table(rows=1, cols=max(1, len(columns)))
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = str(column.get("label") or column.get("key") or "")
    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = str(row.get(column.get("key"), ""))
    data = io.BytesIO()
    doc.save(data)
    data.seek(0)
    return send_file(
        data,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=_safe_filename(review.get("title", "tabella"), "docx"),
    )


def _review_xlsx_response(review: dict[str, Any]):
    columns = review.get("columns") or []
    rows = review.get("rows") or []
    shared_strings: list[str] = []
    string_index: dict[str, int] = {}

    def s(value: Any) -> int:
        text = str(value or "")
        if text not in string_index:
            string_index[text] = len(shared_strings)
            shared_strings.append(text)
        return string_index[text]

    def cell_ref(col: int, row: int) -> str:
        letters = ""
        col += 1
        while col:
            col, rem = divmod(col - 1, 26)
            letters = chr(65 + rem) + letters
        return f"{letters}{row}"

    sheet_rows = []
    header = [s(column.get("label") or column.get("key") or "") for column in columns]
    sheet_rows.append((1, header))
    for row_number, row in enumerate(rows, start=2):
        sheet_rows.append((row_number, [s(row.get(column.get("key"), "")) for column in columns]))
    sheet_data = []
    for row_number, values in sheet_rows:
        cells = "".join(
            f'<c r="{cell_ref(col_index, row_number)}" t="s"><v>{value}</v></c>'
            for col_index, value in enumerate(values)
        )
        sheet_data.append(f'<row r="{row_number}">{cells}</row>')
    shared = "".join(f"<si><t>{html.escape(value)}</t></si>" for value in shared_strings)
    data = io.BytesIO()
    with zipfile.ZipFile(data, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
<Override PartName="/xl/sharedStrings.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sharedStrings+xml"/>
</Types>""")
        zf.writestr("_rels/.rels", """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>""")
        zf.writestr("xl/workbook.xml", """<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<sheets><sheet name="Judicex" sheetId="1" r:id="rId1"/></sheets></workbook>""")
        zf.writestr("xl/_rels/workbook.xml.rels", """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/sharedStrings" Target="sharedStrings.xml"/>
</Relationships>""")
        zf.writestr("xl/worksheets/sheet1.xml", f"""<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"><sheetData>{''.join(sheet_data)}</sheetData></worksheet>""")
        zf.writestr("xl/sharedStrings.xml", f"""<?xml version="1.0" encoding="UTF-8"?>
<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" count="{len(shared_strings)}" uniqueCount="{len(shared_strings)}">{shared}</sst>""")
    data.seek(0)
    return send_file(
        data,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=_safe_filename(review.get("title", "tabella"), "xlsx"),
    )


def _choose_draft_template(instruction: str, custom_templates: list[dict[str, Any]]) -> str:
    text = instruction.lower()
    if "sfratto" in text or "moros" in text or "rilascio" in text:
        return "intimazione_sfratto_morosita"
    if "opposizione" in text or "oppor" in text:
        return "opposizione_decreto_ingiuntivo"
    if "ingiunt" in text or "recupero credito" in text or "credito" in text or "fattura" in text:
        return "ricorso_decreto_ingiuntivo"
    for template in custom_templates:
        haystack = f"{template.get('title', '')} {template.get('name', '')}".lower()
        if any(word and word in text for word in re.findall(r"[a-zà-ù0-9]{4,}", haystack)):
            return str(template.get("id") or template.get("name") or "")
    templates = list_templates()
    return templates[0]["name"] if templates else ""


def _draft_required_params(memory: Any, template_name: str) -> list[str]:
    custom = memory.get_custom_draft_template(template_name)
    if custom is not None:
        return [str(item) for item in custom.get("required_params", [])]
    try:
        template = load_template(template_name)
    except DraftingError:
        return []
    return [str(item) for item in template.get("required_params", [])]


def _draft_template_title(memory: Any, template_name: str) -> str:
    custom = memory.get_custom_draft_template(template_name)
    if custom is not None:
        return str(custom.get("title") or template_name)
    try:
        return str(load_template(template_name).get("title") or template_name)
    except DraftingError:
        return template_name


def _infer_draft_params(
    instruction: str,
    required: list[str],
    *,
    matter: dict[str, Any],
    context: dict[str, Any],
) -> dict[str, str]:
    params: dict[str, str] = {}
    for raw_key, value in _parse_human_fields(instruction).items():
        target = _draft_target_param(raw_key, required)
        if target and value:
            params[target] = value

    client_name = str(matter.get("client_name") or "").strip()
    for target in ("creditore", "attore", "ricorrente", "locatore", "opponente"):
        if target in required and target not in params and client_name:
            params[target] = client_name

    text = re.sub(r"\s+", " ", instruction).strip()
    lowered = text.lower()
    amount = _first_amount(text)
    for target in ("importo", "totale_dovuto", "canone_mensile"):
        if target in required and target not in params and amount:
            params[target] = amount
    if "tribunale" in required and "tribunale" not in params:
        tribunal = _match_first(r"tribunale\s+di\s+([A-Za-zÀ-ÿ' -]{2,60})", text)
        if tribunal:
            params["tribunale"] = tribunal.strip(" .,;")

    counterparty = _match_first(r"(?:contro|nei confronti di|verso)\s+([^.;,\n]{2,80})", text)
    for target in ("debitore", "convenuto", "resistente", "conduttore", "opposto"):
        if target in required and target not in params and counterparty:
            params[target] = counterparty.strip(" .,;")

    if "decreto_numero" in required and "decreto_numero" not in params:
        decreto = _match_first(r"(?:decreto|d\.i\.)\s*(?:n\.?|numero)?\s*([A-Za-z0-9/.-]+)", text)
        if decreto:
            params["decreto_numero"] = decreto.strip(" .,;")
    if "data_notifica" in required and "data_notifica" not in params:
        data_notifica = _match_first(r"notificat[oa]\s+(?:il|in data)?\s*([0-9]{1,2}[/-][0-9]{1,2}[/-][0-9]{2,4})", text)
        if data_notifica:
            params["data_notifica"] = data_notifica.strip(" .,;")

    if "causale" in required and "causale" not in params:
        params["causale"] = _short_instruction_reason(instruction)
    if "motivi_opposizione" in required and "motivi_opposizione" not in params:
        params["motivi_opposizione"] = _short_instruction_reason(instruction)

    parties = context.get("parties") if isinstance(context.get("parties"), list) else []
    amounts = context.get("amounts") if isinstance(context.get("amounts"), list) else []
    if amounts:
        fallback_amount = str(amounts[0].get("value") or amounts[0].get("label") or amounts[0].get("text") or "").strip()
        for target in ("importo", "totale_dovuto"):
            if target in required and target not in params and fallback_amount:
                params[target] = fallback_amount
    if parties:
        fallback_party = str(parties[0].get("label") or parties[0].get("text") or "").strip()
        for target in ("debitore", "convenuto", "conduttore", "opposto"):
            if target in required and target not in params and fallback_party and fallback_party != client_name:
                params[target] = fallback_party

    return {key: value for key, value in params.items() if key in required and str(value).strip()}


def _parse_human_fields(text: str) -> dict[str, str]:
    fields: dict[str, str] = {}
    for line in str(text or "").splitlines():
        clean = line.strip()
        if not clean:
            continue
        separator = ":" if ":" in clean else "=" if "=" in clean else ""
        if not separator:
            continue
        key, _, value = clean.partition(separator)
        key = key.strip()
        value = value.strip()
        if key and value:
            fields[key] = value
    return fields


def _draft_target_param(label: str, required: list[str]) -> str:
    normalized = _normalize_label(label)
    required_set = set(required)
    if normalized in required_set:
        return normalized
    aliases = {
        "creditore": ["creditore", "ricorrente", "attore", "cliente"],
        "debitore": ["debitore", "resistente", "convenuto", "controparte"],
        "opponente": ["opponente", "attore", "cliente"],
        "opposto": ["opposto", "creditore", "controparte"],
        "locatore": ["locatore", "proprietario", "cliente"],
        "conduttore": ["conduttore", "inquilino", "debitore", "controparte"],
        "importo": ["importo", "credito", "somma", "totale"],
        "totale_dovuto": ["totale_dovuto", "totale", "morosita", "importo"],
        "canone_mensile": ["canone_mensile", "canone", "affitto"],
        "mensilita_dovute": ["mensilita_dovute", "mensilita", "mesi"],
        "tribunale": ["tribunale", "foro"],
        "causale": ["causale", "oggetto", "motivo"],
        "motivi_opposizione": ["motivi_opposizione", "motivi", "ragioni", "contestazioni"],
        "immobile": ["immobile", "indirizzo", "locale"],
        "decreto_numero": ["decreto_numero", "decreto", "numero_decreto"],
        "data_notifica": ["data_notifica", "notifica", "data"],
    }
    for target, names in aliases.items():
        if target in required_set and normalized in {_normalize_label(item) for item in names}:
            return target
    return ""


def _normalize_label(value: str) -> str:
    text = str(value or "").strip().lower()
    replacements = str.maketrans("àèéìòù", "aeeiou")
    text = text.translate(replacements)
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_")


def _first_amount(text: str) -> str:
    patterns = [
        r"(?:€|euro)\s*([0-9][0-9.,]*)",
        r"([0-9][0-9.,]*)\s*(?:€|euro)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return ""


def _match_first(pattern: str, text: str) -> str:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return match.group(1) if match else ""


def _short_instruction_reason(instruction: str) -> str:
    text = re.sub(r"\s+", " ", instruction).strip()
    return text[:600]


def _friendly_draft_error(message: str) -> str:
    if "missing required params" in message:
        missing = re.findall(r"'([^']+)'", message)
        if missing:
            return "Mi servono questi dati: " + ", ".join(missing)
    return message


def _render_custom_template(template: dict[str, Any], params: dict[str, str]) -> str:
    body = str(template.get("body") or "")
    required = [str(item) for item in template.get("required_params", [])]
    missing = [key for key in required if not params.get(key)]
    if missing:
        raise DraftingError(f"parametri mancanti: {missing}")

    def replace(match: re.Match[str]) -> str:
        key = match.group(1)
        return str(params.get(key, match.group(0)))

    title = str(template.get("title") or "").strip()
    rendered = re.sub(r"\{([A-Za-z_][A-Za-z0-9_]*)\}", replace, body)
    return f"{title.upper()}\n\n{rendered}".strip() if title else rendered.strip()


def _is_auth_open_path(path: str) -> bool:
    if path.startswith("/static/"):
        return True
    return path in {
        "/api/auth/status",
        "/api/auth/login",
        "/api/auth/logout",
        "/api/auth/setup",
    }


def _local_auth_state(app: Flask, store_factory: Any) -> dict[str, Any]:
    env_hash = str(app.config.get("JUDICEX_LOCAL_PASSWORD_HASH", "")).strip()
    if env_hash:
        return {"configured": True, "password_hash": env_hash, "source": "env_hash"}
    env_password = str(app.config.get("JUDICEX_LOCAL_PASSWORD", "")).strip()
    if env_password:
        return {"configured": True, "plain_password": env_password, "source": "env"}
    try:
        with store_factory() as memory:
            setting = memory.get_app_setting("local_auth") or {}
    except Exception:
        setting = {}
    password_hash = str(setting.get("password_hash", "")).strip()
    enabled = bool(setting.get("enabled", bool(password_hash)))
    return {
        "configured": bool(enabled and password_hash),
        "password_hash": password_hash,
        "source": "sqlite" if password_hash else "",
    }


def _hash_password(password: str) -> str:
    salt = os.urandom(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 240_000)
    return "pbkdf2_sha256$240000$" + base64.b64encode(salt).decode("ascii") + "$" + base64.b64encode(digest).decode("ascii")


def _verify_local_password(password: str, auth: dict[str, Any]) -> bool:
    plain = auth.get("plain_password")
    if plain:
        return hmac.compare_digest(str(plain), password)
    encoded = str(auth.get("password_hash") or "")
    try:
        algorithm, iterations_raw, salt_raw, digest_raw = encoded.split("$", 3)
        if algorithm != "pbkdf2_sha256":
            return False
        salt = base64.b64decode(salt_raw.encode("ascii"))
        expected = base64.b64decode(digest_raw.encode("ascii"))
        digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, int(iterations_raw))
        return hmac.compare_digest(digest, expected)
    except Exception:
        return False


def _login_page() -> str:
    return """
<!doctype html>
<html lang="it">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Judicex - Accesso</title>
    <style>
      body{margin:0;min-height:100vh;display:grid;place-items:center;font-family:Inter,system-ui,sans-serif;color:#111;background:#fff}
      main{width:min(420px,calc(100vw - 32px));border:1px solid #e5e5e5;border-radius:10px;padding:24px;box-shadow:0 20px 60px rgba(0,0,0,.08)}
      h1{font-family:Georgia,serif;font-weight:400;margin:0 0 8px;font-size:32px}
      p{margin:0 0 18px;color:#666;line-height:1.5}
      input,button{width:100%;height:42px;border-radius:8px;font:inherit;box-sizing:border-box}
      input{border:1px solid #d4d4d8;padding:0 12px}
      button{margin-top:12px;border:0;background:#111;color:#fff;font-weight:600}
      .error{min-height:20px;margin-top:12px;color:#b91c1c;font-size:14px}
    </style>
  </head>
  <body>
    <main>
      <h1>Judicex</h1>
      <p>Inserisci la password locale per aprire la memoria di lavoro.</p>
      <form id="loginForm">
        <input id="password" type="password" autocomplete="current-password" placeholder="Password locale" autofocus>
        <button>Entra</button>
        <div id="error" class="error"></div>
      </form>
    </main>
    <script>
      document.getElementById("loginForm").addEventListener("submit", async function(event){
        event.preventDefault();
        const error = document.getElementById("error");
        error.textContent = "";
        const response = await fetch("/api/auth/login", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({password: document.getElementById("password").value})
        });
        if (response.ok) window.location.reload();
        else error.textContent = (await response.json().catch(() => ({}))).error || "Accesso non riuscito.";
      });
    </script>
  </body>
</html>
""".strip()


def _backup_response(db_path: Path):
    db_path = db_path.resolve()
    files_dir = db_path.parent / f"{db_path.stem}_files"
    data = io.BytesIO()
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_db = Path(temp_dir) / "memory.db"
        with LegalMemoryStore(db_path) as memory:
            backup_conn = sqlite3.connect(str(temp_db))
            try:
                memory.conn.backup(backup_conn)
            finally:
                backup_conn.close()
        with zipfile.ZipFile(data, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr(
                "judicex_backup.json",
                json.dumps(
                    {
                        "app": "judicex",
                        "created_at": datetime.now(timezone.utc).isoformat(),
                        "database": "memory.db",
                        "files_dir": "files",
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
            )
            archive.write(temp_db, "memory.db")
            if files_dir.is_dir():
                for path in files_dir.rglob("*"):
                    if path.is_file():
                        archive.write(path, Path("files") / path.relative_to(files_dir))
    data.seek(0)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    return send_file(
        data,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"judicex_backup_{stamp}.zip",
    )


def _restore_backup(db_path: Path, upload: Any) -> dict[str, Any]:
    db_path = db_path.resolve()
    db_path.parent.mkdir(parents=True, exist_ok=True)
    files_dir = db_path.parent / f"{db_path.stem}_files"
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_dir_path = Path(temp_dir)
        backup_path = temp_dir_path / "backup.zip"
        upload.save(backup_path)
        with zipfile.ZipFile(backup_path) as archive:
            names = set(archive.namelist())
            if "memory.db" not in names or "judicex_backup.json" not in names:
                raise ValueError("Backup Judicex non valido.")
            restored_db = temp_dir_path / "memory.db"
            archive.extract("memory.db", temp_dir_path)
            with LegalMemoryStore(restored_db) as restored:
                restored.health()
            if db_path.exists():
                shutil.copy2(db_path, db_path.with_name(f"{db_path.stem}.pre_restore_{stamp}{db_path.suffix}"))
            for suffix in ("-wal", "-shm"):
                sidecar = Path(str(db_path) + suffix)
                if sidecar.exists():
                    sidecar.unlink()
            shutil.copy2(restored_db, db_path)
            file_members = [name for name in archive.namelist() if name.startswith("files/") and not name.endswith("/")]
            if file_members:
                if files_dir.exists():
                    archived_files = files_dir.with_name(f"{files_dir.name}.pre_restore_{stamp}")
                    if archived_files.exists():
                        shutil.rmtree(archived_files)
                    shutil.move(str(files_dir), str(archived_files))
                files_dir.mkdir(parents=True, exist_ok=True)
                for name in file_members:
                    target = files_dir / Path(name).relative_to("files")
                    target.parent.mkdir(parents=True, exist_ok=True)
                    with archive.open(name) as source, target.open("wb") as destination:
                        shutil.copyfileobj(source, destination)
    return {"status": "restored", "db": str(db_path)}


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Judicex Flask web UI.")
    parser.add_argument("--db", required=True, help="Path to the Judicex SQLite database.")
    parser.add_argument("--model", default="", help="Default model fallback. The UI provider settings can override it.")
    parser.add_argument("--area", default="civile", help="Default legal area.")
    parser.add_argument("--host", default="http://127.0.0.1:11434", help="Legacy Ollama host fallback.")
    parser.add_argument("--bind", default="127.0.0.1", help="Flask bind address.")
    parser.add_argument("--port", type=int, default=5050, help="Flask port.")
    parser.add_argument("--debug", action="store_true", help="Run Flask in debug mode.")
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    app = create_app(
        db_path=args.db,
        default_model=args.model,
        default_area=args.area,
        ollama_host=args.host,
    )
    app.run(host=args.bind, port=args.port, debug=args.debug)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
